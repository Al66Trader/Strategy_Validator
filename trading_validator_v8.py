# ==============================================================
# TRADING STRATEGY VALIDATOR v8
# Framework Marco Morrone + CSCV Bailey & De Prado (2015)
# Archivio locale strategie
#
# Formati supportati:
#   Tab 1 (Morrone): CSV standard  OPPURE  XLS MultiCharts
#   Tab 2 (CSCV):    N file CSV (uno per configurazione)
#   Tab 3 (Archivio): salvataggio e rielaborazione risultati
#
# Installazione:
#   pip install streamlit pandas numpy scipy matplotlib openpyxl
# Avvio:
#   streamlit run trading_validator_v8.py
# ==============================================================


import streamlit as st
import pandas as pd
import numpy as np
from scipy import stats
from scipy.optimize import brentq
from itertools import combinations
from math import comb
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
import io, json, shutil
from datetime import datetime
from pathlib import Path
import warnings
warnings.filterwarnings("ignore")

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


# ══════════════════════════════════════════════════════════════
# MODULO ARCHIVIO
# Cartella "validator_archive" nella stessa dir del .py
# ══════════════════════════════════════════════════════════════
ARCHIVE_DIR   = Path(__file__).parent / "validator_archive"
MANIFEST_FILE = ARCHIVE_DIR / "manifest.json"


def _sv(v):
    if isinstance(v, (np.floating, np.integer)):
        return float(v)
    if isinstance(v, (int, float, bool, str)):
        return v
    return str(v)


def arc_init():
    ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)
    if not MANIFEST_FILE.exists():
        MANIFEST_FILE.write_text("[]", encoding="utf-8")


def arc_manifest():
    arc_init()
    try:
        d = json.loads(MANIFEST_FILE.read_text(encoding="utf-8"))
        return d if isinstance(d, list) else []
    except Exception:
        return []


def arc_save_manifest(m):
    MANIFEST_FILE.write_text(
        json.dumps(m, indent=2, ensure_ascii=False), encoding="utf-8"
    )


def arc_make_id(name):
    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = "".join(c if c.isalnum() or c in "-_" else "_" for c in name)
    return f"{safe[:30].strip('_')}_{ts}"


def arc_save_morrone(strategy_name, instrument, initial_capital,
                     is_end_date, n_totale, alpha, dsr_threshold,
                     trades_df, m_is, m_oos, test_results, n_passed, note=""):
    arc_init()
    eid  = arc_make_id(strategy_name)
    edir = ARCHIVE_DIR / eid
    edir.mkdir(parents=True, exist_ok=True)
    trades_df.to_csv(edir / "trades.csv", index=False)

    tests_clean = {
        k: {fk: _sv(fv) for fk, fv in v.items() if fk != "passed"}
           | {"passed": bool(v["passed"])}
        for k, v in test_results.items()
    }

    def cm(m):
        return {k: _sv(v) for k, v in m.items()
                if k not in ("returns", "profits")} if m else None

    (edir / "morrone_results.json").write_text(
        json.dumps({"n_passed": n_passed, "tests": tests_clean,
                    "metrics_is": cm(m_is), "metrics_oos": cm(m_oos)},
                   indent=2, ensure_ascii=False), encoding="utf-8"
    )

    vmap = {5: "VALIDA", 4: "QUASI VALIDA", 3: "PARZIALE",
            2: "NON IDONEA", 1: "NON IDONEA", 0: "NON IDONEA"}
    meta = {
        "id": eid, "strategy_name": strategy_name,
        "instrument": instrument, "initial_capital": initial_capital,
        "is_end_date": str(is_end_date), "n_totale": n_totale,
        "alpha": alpha, "dsr_threshold": dsr_threshold,
        "archived_at": datetime.now().isoformat(),
        "has_morrone": True, "has_cscv": False,
        "verdetto": vmap.get(n_passed, "NON IDONEA"),
        "n_passed": n_passed, "pbo": None, "note": note,
    }
    (edir / "metadata.json").write_text(
        json.dumps(meta, indent=2, ensure_ascii=False), encoding="utf-8"
    )

    manifest = arc_manifest()
    manifest.insert(0, {
        "id": eid, "strategy_name": strategy_name,
        "instrument": instrument,
        "archived_at": datetime.now().isoformat(),
        "verdetto": vmap.get(n_passed, "NON IDONEA"),
        "n_passed": n_passed, "has_cscv": False, "pbo": None, "note": note,
    })
    arc_save_manifest(manifest)
    return eid


def arc_add_cscv(eid, cscv_result, matrix_df):
    edir = ARCHIVE_DIR / eid
    if not edir.exists():
        return False
    matrix_df.to_csv(edir / "monthly_pnl.csv")
    (edir / "cscv_results.json").write_text(json.dumps({
        "pbo": float(cscv_result["pbo"]),
        "prob_loss": float(cscv_result["prob_loss"]),
        "N_configs": int(cscv_result["N_configs"]),
        "T_months": int(cscv_result["T_months"]),
        "S": int(cscv_result["S"]),
        "n_combos": int(cscv_result["n_combos"]),
        "logits":      cscv_result["logits"].tolist(),
        "sr_is_list":  cscv_result["sr_is_list"].tolist(),
        "sr_oos_list": cscv_result["sr_oos_list"].tolist(),
        "neff":        int(cscv_result.get("neff", 0)) if cscv_result.get("neff") else None,
        "n_nominal":   int(cscv_result.get("n_nominal", 0)) if cscv_result.get("n_nominal") else None,
        "verdetto_neff": cscv_result.get("verdetto_neff"),
        "n_pass_neff":   int(cscv_result.get("n_pass_neff", 0)) if cscv_result.get("n_pass_neff") is not None else None,
    }, indent=2), encoding="utf-8")
    mp   = edir / "metadata.json"
    meta = json.loads(mp.read_text(encoding="utf-8"))
    meta["has_cscv"] = True
    meta["pbo"]      = float(cscv_result["pbo"])
    mp.write_text(json.dumps(meta, indent=2, ensure_ascii=False), encoding="utf-8")
    manifest = arc_manifest()
    for item in manifest:
        if item["id"] == eid:
            item["has_cscv"] = True
            item["pbo"]      = float(cscv_result["pbo"])
    arc_save_manifest(manifest)
    return True


def arc_load(eid):
    edir = ARCHIVE_DIR / eid
    if not edir.exists():
        return None, "Entry non trovata"
    result = {}
    try:
        result["metadata"] = json.loads(
            (edir / "metadata.json").read_text(encoding="utf-8")
        )
    except Exception as e:
        return None, str(e)
    for fname, key, kwargs in [
        ("trades.csv",      "trades",      {"parse_dates": ["entry_datetime","exit_datetime"]}),
        ("monthly_pnl.csv", "monthly_pnl", {"index_col": 0}),
    ]:
        p = edir / fname
        if p.exists():
            try:
                result[key] = pd.read_csv(p, **kwargs)
            except Exception:
                result[key] = None
    for fname, key in [("morrone_results.json","morrone"), ("cscv_results.json","cscv")]:
        p = edir / fname
        if p.exists():
            try:
                raw = json.loads(p.read_text(encoding="utf-8"))
                if key == "cscv":
                    for k in ("logits","sr_is_list","sr_oos_list"):
                        raw[k] = np.array(raw[k])
                result[key] = raw
            except Exception:
                result[key] = None
    return result, None


def arc_delete(eid):
    edir = ARCHIVE_DIR / eid
    if edir.exists():
        shutil.rmtree(edir)
    m = arc_manifest()
    arc_save_manifest([x for x in m if x["id"] != eid])


arc_init()




# ══════════════════════════════════════════════════════════════
# EXPORT WORD — Genera report .docx leggibile da chiunque
# ══════════════════════════════════════════════════════════════
def _set_cell_bg(cell, hex_color):
    """Colore sfondo cella Word."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _bold_cell(cell, text, size=10, color=None, bg=None,
               align=WD_ALIGN_PARAGRAPH.LEFT):
    if bg:
        _set_cell_bg(cell, bg)
    p   = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _normal_cell(cell, text, size=10, bold=False,
                 align=WD_ALIGN_PARAGRAPH.LEFT):
    p   = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold      = bold
    run.font.size = Pt(size)


def arc_export_docx(entry):
    """
    Genera un documento Word (.docx) con il report completo
    di una validazione archiviata. Restituisce (bytes, None)
    oppure (None, messaggio_errore).
    """
    if not DOCX_OK:
        return None, ("Libreria python-docx non installata. "
                      "Riavvia il .bat per installarla automaticamente.")
    try:
        meta  = entry.get("metadata", {})
        mor   = entry.get("morrone",  {})
        cscv  = entry.get("cscv",     None)
        m_is  = mor.get("metrics_is",  {}) or {}
        m_oos = mor.get("metrics_oos", {}) or {}
        tests = mor.get("tests",       {}) or {}
        n_p   = meta.get("n_passed", 0)
        verd  = meta.get("verdetto",
                         {5:"VALIDA", 4:"QUASI VALIDA", 3:"PARZIALE",
                          2:"NON IDONEA", 1:"NON IDONEA",
                          0:"NON IDONEA"}.get(n_p, "NON IDONEA"))

        doc = DocxDocument()

        # Margini pagina
        for sec in doc.sections:
            sec.top_margin    = Cm(2)
            sec.bottom_margin = Cm(2)
            sec.left_margin   = Cm(2.5)
            sec.right_margin  = Cm(2.5)

        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)

        # ── Titolo ───────────────────────────────────────────
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h.add_run("TRADING STRATEGY VALIDATOR v8")
        r.bold = True
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1A, 0x2B, 0x5E)

        h2 = doc.add_paragraph()
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = h2.add_run("Report di Validazione — Framework Morrone + CSCV")
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        doc.add_paragraph()

        # ── Verdetto ─────────────────────────────────────────
        color_v = ("02A87A" if verd == "VALIDA"
                   else ("F59E0B" if "PARZIALE" in verd else "E84D4D"))
        vp = doc.add_paragraph()
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = vp.add_run(f"VERDETTO:  {verd}  ({n_p}/5 test superati)")
        rv.bold = True
        rv.font.size = Pt(14)
        rv.font.color.rgb = RGBColor.from_string(color_v)
        doc.add_paragraph()

        # ── Dati strategia ───────────────────────────────────
        doc.add_heading("Dati Strategia", level=1)
        t1 = doc.add_table(rows=2, cols=4)
        t1.style = "Table Grid"
        hdrs1 = ["Strategia", "Strumento", "Capitale Iniziale", "Data Fine IS"]
        vals1 = [meta.get("strategy_name", "—"),
                 meta.get("instrument",    "—"),
                 f"${meta.get('initial_capital', 0):,}",
                 meta.get("is_end_date", "—")]
        for i, (h_txt, v_txt) in enumerate(zip(hdrs1, vals1)):
            _bold_cell(t1.cell(0, i), h_txt, bg="1A2B5E", color="FFFFFF",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _normal_cell(t1.cell(1, i), v_txt,
                         align=WD_ALIGN_PARAGRAPH.CENTER)

        doc.add_paragraph()
        t2 = doc.add_table(rows=2, cols=3)
        t2.style = "Table Grid"
        hdrs2 = ["N Combinazioni testate", "Archiviato il", "Note"]
        vals2 = [f"{meta.get('n_totale', 0):,}",
                 meta.get("archived_at", "—")[:10],
                 meta.get("note", "—") or "—"]
        for i, (h_txt, v_txt) in enumerate(zip(hdrs2, vals2)):
            _bold_cell(t2.cell(0, i), h_txt, bg="1A2B5E", color="FFFFFF",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _normal_cell(t2.cell(1, i), v_txt,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

        # ── Test statistici ──────────────────────────────────
        doc.add_heading("Test Statistici (Framework Morrone)", level=1)
        if tests:
            t3 = doc.add_table(rows=len(tests) + 1, cols=4)
            t3.style = "Table Grid"
            for j, h_txt in enumerate(["Test", "Valore", "Soglia", "Esito"]):
                _bold_cell(t3.cell(0, j), h_txt, bg="2E5BBA", color="FFFFFF",
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            for i, (tk, tv) in enumerate(tests.items(), start=1):
                passed = tv.get("passed", False)
                _normal_cell(t3.cell(i, 0), tv.get("name", tk))
                _normal_cell(t3.cell(i, 1), str(tv.get("value",  "—")),
                             align=WD_ALIGN_PARAGRAPH.CENTER)
                _normal_cell(t3.cell(i, 2), str(tv.get("soglia", "—")),
                             align=WD_ALIGN_PARAGRAPH.CENTER)
                c3 = t3.cell(i, 3)
                _set_cell_bg(c3, "E8F5E9" if passed else "FFEBEE")
                p3 = c3.paragraphs[0]
                p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r3 = p3.add_run("PASS" if passed else "FAIL")
                r3.bold = True
                r3.font.size = Pt(10)
                r3.font.color.rgb = RGBColor.from_string(
                    "02A87A" if passed else "E84D4D")
        doc.add_paragraph()

        # ── Metriche IS / OOS ────────────────────────────────
        doc.add_heading("Metriche Performance", level=1)
        has_oos = bool(m_oos)
        metr_rows = [
            ("Trade",
             str(int(m_is.get("n_trades", 0))),
             str(int(m_oos.get("n_trades", 0))) if has_oos else "—"),
            ("Sharpe Ratio",
             f"{float(m_is.get('sharpe', 0)):.3f}",
             f"{float(m_oos.get('sharpe', 0)):.3f}" if has_oos else "—"),
            ("CAGR",
             f"{float(m_is.get('cagr', 0)):.2%}",
             f"{float(m_oos.get('cagr', 0)):.2%}" if has_oos else "—"),
            ("Max Drawdown",
             f"{float(m_is.get('max_dd', 0)):.2%}",
             f"{float(m_oos.get('max_dd', 0)):.2%}" if has_oos else "—"),
            ("Profit Factor",
             f"{float(m_is.get('profit_factor', 0)):.2f}",
             f"{float(m_oos.get('profit_factor', 0)):.2f}" if has_oos else "—"),
            ("Win Rate",
             f"{float(m_is.get('win_rate', 0)):.1%}",
             f"{float(m_oos.get('win_rate', 0)):.1%}" if has_oos else "—"),
            ("Avg Trade ($)",
             f"${float(m_is.get('avg_trade', 0)):.0f}",
             f"${float(m_oos.get('avg_trade', 0)):.0f}" if has_oos else "—"),
        ]
        t4 = doc.add_table(rows=len(metr_rows) + 1, cols=3)
        t4.style = "Table Grid"
        for j, h_txt in enumerate(["Metrica", "In-Sample", "Out-of-Sample"]):
            _bold_cell(t4.cell(0, j), h_txt, bg="2E5BBA", color="FFFFFF",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for i, (label, is_v, oos_v) in enumerate(metr_rows, start=1):
            bg_r = "F8FAFC" if i % 2 == 0 else "FFFFFF"
            for j in range(3):
                _set_cell_bg(t4.cell(i, j), bg_r)
            _normal_cell(t4.cell(i, 0), label, bold=True)
            _normal_cell(t4.cell(i, 1), is_v,  align=WD_ALIGN_PARAGRAPH.CENTER)
            _normal_cell(t4.cell(i, 2), oos_v, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

        # ── CSCV ─────────────────────────────────────────────
        if cscv:
            doc.add_heading(
                "CSCV — Combinatorially Symmetric Cross-Validation", level=1)
            pr = doc.add_paragraph()
            rr = pr.add_run(
                "Bailey, Borwein, Lopez de Prado, Zhu — "
                "J. Computational Finance, 2015")
            rr.italic = True
            rr.font.size = Pt(9)

            pbo = float(cscv.get("pbo", 0))
            pl  = float(cscv.get("prob_loss", 0))
            esito_c = ("ROBUSTO  (PBO < 5%)" if pbo < 0.05
                       else ("ATTENZIONE  (5-15%)" if pbo < 0.15
                             else "OVERFITTING  (> 15%)"))
            color_c = ("02A87A" if pbo < 0.05
                       else ("F59E0B" if pbo < 0.15 else "E84D4D"))
            pc = doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rc = pc.add_run(f"PBO = {pbo:.1%}  —  {esito_c}")
            rc.bold = True
            rc.font.size = Pt(13)
            rc.font.color.rgb = RGBColor.from_string(color_c)

            cscv_rows = [
                ("PBO (Prob. Backtest Overfitting)",
                 f"{pbo:.1%}", "Obiettivo < 5%"),
                ("Probabilita perdita OOS",
                 f"{pl:.1%}", "—"),
                ("N configurazioni testate",
                 str(cscv.get("N_configs", "—")), "Min 4"),
                ("N combinazioni IS/OOS calcolate",
                 f"{int(cscv.get('n_combos', 0)):,}", "—"),
                ("Mesi totali (T)",
                 str(cscv.get("T_months", "—")), "—"),
                ("Partizioni (S)",
                 str(cscv.get("S", "—")), "Consigliato 16"),
            ]
            t5 = doc.add_table(rows=len(cscv_rows) + 1, cols=3)
            t5.style = "Table Grid"
            for j, h_txt in enumerate(["Parametro", "Valore", "Soglia"]):
                _bold_cell(t5.cell(0, j), h_txt, bg="2E5BBA", color="FFFFFF",
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            for i, (label, val, soglia) in enumerate(cscv_rows, start=1):
                bg_r = "F8FAFC" if i % 2 == 0 else "FFFFFF"
                for j in range(3):
                    _set_cell_bg(t5.cell(i, j), bg_r)
                _normal_cell(t5.cell(i, 0), label, bold=True)
                _normal_cell(t5.cell(i, 1), val,
                             align=WD_ALIGN_PARAGRAPH.CENTER)
                _normal_cell(t5.cell(i, 2), soglia,
                             align=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_paragraph()

        # ── Guida lettura ────────────────────────────────────
        doc.add_heading("Legenda — Test Statistici (Livello 1)", level=1)
        p_leg = doc.add_paragraph()
        p_leg.add_run(
            "Ogni test valuta un aspetto diverso della robustezza statistica. "
            "La strategia supera il livello 1 se ottiene almeno 5/5 test."
        ).font.size = Pt(9)
        doc.add_paragraph()

        # Tabella legenda a 4 colonne
        legend_l1 = [
            ("T-Statistic",
             "Verifica che il profitto medio per trade sia statisticamente "
             "diverso da zero e non frutto del caso.",
             "t > 1.96 (con alpha=5%). "
             "Piu alto e', piu il profitto e' solido statisticamente.",
             "PASS se t > 1.96\nFAIL se t <= 1.96"),
            ("E[MaxSR] / EMSR",
             "Confronta lo Sharpe Ratio della strategia con quello atteso "
             "da N strategie casuali. Misura se il risultato e' davvero "
             "buono o solo il migliore per caso tra tanti tentativi.",
             "SR osservato > E[MaxSR]. "
             "Con N=367 ottimizzazioni il benchmark casuale e' circa 0.795.",
             "PASS se SR > E[MaxSR]\nFAIL se SR <= E[MaxSR]"),
            ("MinBTL",
             "Anni minimi di backtest necessari perche' i risultati siano "
             "statisticamente affidabili, dato il numero di ottimizzazioni N.",
             "Anni IS >= MinBTL. "
             "Con piu' ottimizzazioni serve un backtest piu' lungo "
             "per avere la stessa confidenza statistica.",
             "PASS se Anni IS >= MinBTL\nFAIL se Anni IS < MinBTL"),
            ("Bonferroni",
             "Correzione per il problema dei test multipli: testare N "
             "configurazioni aumenta la probabilita' di trovare un risultato "
             "positivo per puro caso. Alza la soglia di significativita'.",
             "t > soglia Bonferroni = t_critico(alpha/N). "
             "Piu' alto e' N, piu' alta deve essere la t-stat.",
             "PASS se t > t_bonf\nFAIL se t <= t_bonf"),
            ("DSR — Deflated Sharpe Ratio",
             "Versione corretta dello Sharpe che tiene conto di: numero di "
             "ottimizzazioni N, non-normalita' dei ritorni (skewness e "
             "kurtosis), lunghezza del backtest. E' il test piu' completo.",
             "DSR > soglia (default 0.95). "
             "Misura la probabilita' che lo Sharpe osservato sia reale "
             "e non inflazionato dall'ottimizzazione.",
             "PASS se DSR > 0.95\nFAIL se DSR <= 0.95"),
        ]

        t6 = doc.add_table(rows=len(legend_l1) + 1, cols=4)
        t6.style = "Table Grid"
        hdrs6 = ["Test", "Cosa misura", "Come si legge", "PASS / FAIL"]
        bg_hdrs = ["1A2B5E", "1A2B5E", "1A2B5E", "1A2B5E"]
        for j, h_txt in enumerate(hdrs6):
            _bold_cell(t6.cell(0, j), h_txt, bg=bg_hdrs[j],
                       color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        for i, (test, cosa, come, esito) in enumerate(legend_l1, start=1):
            bg_r = "F8FAFC" if i % 2 == 0 else "FFFFFF"
            for j in range(4):
                _set_cell_bg(t6.cell(i, j), bg_r)
            _normal_cell(t6.cell(i, 0), test, bold=True)
            _normal_cell(t6.cell(i, 1), cosa)
            _normal_cell(t6.cell(i, 2), come)
            # Colonna PASS/FAIL con colore
            c3 = t6.cell(i, 3)
            p3 = c3.paragraphs[0]
            for line in esito.split("\n"):
                is_pass = line.startswith("PASS")
                r3 = p3.add_run(line)
                r3.bold = True
                r3.font.size = Pt(9)
                r3.font.color.rgb = RGBColor.from_string(
                    "02A87A" if is_pass else "E84D4D")
                if line != esito.split("\n")[-1]:
                    p3.add_run("\n")
        doc.add_paragraph()

        # Nota verdetto finale
        p_verd = doc.add_paragraph()
        rv2 = p_verd.add_run(
            "Verdetto finale:  5/5 = VALIDA  |  4/5 = QUASI VALIDA  |  "
            "3/5 = PARZIALE  |  2/5 o meno = NON IDONEA"
        )
        rv2.bold = True; rv2.font.size = Pt(9)
        rv2.font.color.rgb = RGBColor(0x1A, 0x2B, 0x5E)
        doc.add_paragraph()

        # ── Footer ───────────────────────────────────────────
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf = fp.add_run(
            f"Generato da Trading Strategy Validator v8  |  "
            f"{datetime.now().strftime('%d/%m/%Y %H:%M')}  |  "
            "Bailey & De Prado (2015)  -  Harvey & Liu (2014)  -  Marco Morrone"
        )
        rf.font.size = Pt(8)
        rf.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue(), None

    except Exception as e:
        import traceback
        return None, f"{e}\n{traceback.format_exc()}"


def cscv_export_docx(cscv_result, interp_lines, strategy_name, img_bytes):
    """
    Genera un documento Word (.docx) con il report completo CSCV.
    Restituisce (bytes, None) oppure (None, messaggio_errore).
    """
    if not DOCX_OK:
        return None, ("Libreria python-docx non installata. "
                      "Riavvia il .bat per installarla automaticamente.")
    try:
        pbo = float(cscv_result.get("pbo", 0))
        pl  = float(cscv_result.get("prob_loss", 0))
        n_configs  = cscv_result.get("N_configs", "—")
        n_combos   = cscv_result.get("n_combos",  0)
        t_months   = cscv_result.get("T_months",  "—")
        s_val      = cscv_result.get("S",          "—")

        esito_c = ("ROBUSTO  (PBO < 5%)" if pbo < 0.05
                   else ("ATTENZIONE  (5–15%)" if pbo < 0.15
                         else "OVERFITTING  (> 15%)"))
        color_c = ("02A87A" if pbo < 0.05
                   else ("F59E0B" if pbo < 0.15 else "E84D4D"))

        doc = DocxDocument()
        for sec in doc.sections:
            sec.top_margin    = Cm(2)
            sec.bottom_margin = Cm(2)
            sec.left_margin   = Cm(2.5)
            sec.right_margin  = Cm(2.5)
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)

        # ── Titolo ───────────────────────────────────────────
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h.add_run("TRADING STRATEGY VALIDATOR v8")
        r.bold = True; r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1A, 0x2B, 0x5E)

        h2 = doc.add_paragraph()
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = h2.add_run(f"Report CSCV — {strategy_name}")
        r2.font.size = Pt(12)
        r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        ref = doc.add_paragraph()
        ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = ref.add_run(
            "Bailey, Borwein, Lopez de Prado, Zhu — "
            "J. Computational Finance, 2015")
        rr.italic = True; rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        doc.add_paragraph()

        # ── Verdetto CSCV ────────────────────────────────────
        vp = doc.add_paragraph()
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rv = vp.add_run(f"PBO = {pbo:.1%}  —  {esito_c}")
        rv.bold = True; rv.font.size = Pt(15)
        rv.font.color.rgb = RGBColor.from_string(color_c)
        doc.add_paragraph()

        # ── Metriche principali ──────────────────────────────
        doc.add_heading("Risultati CSCV", level=1)
        cscv_rows = [
            ("PBO — Probability of Backtest Overfitting",
             f"{pbo:.1%}", "Obiettivo < 5%"),
            ("Probabilita perdita OOS",
             f"{pl:.1%}", "—"),
            ("N configurazioni testate",
             str(n_configs), "Minimo 4"),
            ("N combinazioni IS/OOS calcolate",
             f"{int(n_combos):,}", "—"),
            ("Mesi totali (T)",
             str(t_months), "—"),
            ("Partizioni (S)",
             str(s_val), "Consigliato 16"),
        ]
        t1 = doc.add_table(rows=len(cscv_rows)+1, cols=3)
        t1.style = "Table Grid"
        for j, h_txt in enumerate(["Parametro", "Valore", "Soglia"]):
            _bold_cell(t1.cell(0, j), h_txt, bg="2E5BBA", color="FFFFFF",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for i, (label, val, soglia) in enumerate(cscv_rows, start=1):
            bg_r = "F8FAFC" if i % 2 == 0 else "FFFFFF"
            for j in range(3):
                _set_cell_bg(t1.cell(i, j), bg_r)
            _normal_cell(t1.cell(i, 0), label, bold=True)
            _normal_cell(t1.cell(i, 1), val,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
            _normal_cell(t1.cell(i, 2), soglia,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

        # ── Report testuale ──────────────────────────────────
        doc.add_heading("Report dettagliato", level=1)
        interp_text = "\n".join(interp_lines)
        p_rep = doc.add_paragraph()
        p_rep.style = doc.styles["No Spacing"] if "No Spacing" in             [s.name for s in doc.styles] else doc.styles["Normal"]
        run_rep = p_rep.add_run(interp_text)
        run_rep.font.name = "Courier New"
        run_rep.font.size = Pt(9)
        doc.add_paragraph()

        # ── Grafico ──────────────────────────────────────────
        doc.add_heading("Grafici CSCV", level=1)
        p_gr = doc.add_paragraph()
        p_gr.add_run(
            "Distribuzione dei logit IS vs OOS e scatter Sharpe IS/OOS "
            "su tutte le combinazioni calcolate."
        ).font.size = Pt(9)
        img_stream = io.BytesIO(img_bytes)
        doc.add_picture(img_stream, width=Cm(15))
        doc.add_paragraph()

        # ── Guida lettura CSCV ───────────────────────────────
        doc.add_heading("Legenda — Parametri CSCV (Livello 2)", level=1)
        p_leg2 = doc.add_paragraph()
        run_leg2 = p_leg2.add_run(
            "La CSCV valuta se il processo di selezione dei parametri "
            "e' affetto da overfitting. Opera a livello di scelta dei parametri "
            "all'interno di una singola strategia, non a livello di portafoglio."
        )
        run_leg2.font.size = Pt(9)
        doc.add_paragraph()

        legend_l2 = [
            ("PBO",
             "Probability of Backtest Overfitting. Misura quante volte, "
             "su tutte le combinazioni IS/OOS simmetriche, "
             "il parametro migliore IS risulta peggiore del benchmark OOS.",
             "PBO < 5%: selezione robusta | PBO 5-15%: attenzione | PBO > 15%: overfitting",
             "PASS se PBO < 5% | FAIL se PBO >= 5%"),
            ("Probabilita perdita OOS",
             "Percentuale di sotto-periodi OOS in cui la strategia selezionata IS "
             "produce un rendimento negativo. Misura il rischio concreto di perdita.",
             "Piu bassa e' meglio. Valori > 50% indicano perdita nella maggioranza dei periodi OOS.",
             "Riferimento: < 30%"),
            ("Beta degradazione",
             "Pendenza della regressione SR-IS vs SR-OOS. "
             "Indica se uno Sharpe alto IS predice uno Sharpe alto o basso OOS.",
             "Beta > 0: SR alto IS -> SR alto OOS (buono) | "
             "Beta = 0: nessuna relazione | "
             "Beta < 0: SR alto IS -> SR basso OOS (overfitting classico)",
             "PASS indicativo se Beta > 0 | FAIL indicativo se Beta < 0"),
            ("R2 regressione",
             "Coefficiente di determinazione della regressione SR-IS vs SR-OOS. "
             "Misura quanta parte della varianza OOS e' spiegata dalla performance IS.",
             "R2 vicino a 1: forte relazione IS/OOS (buono) | "
             "R2 vicino a 0: performance IS e OOS non correlate",
             "Alto = buono. Indicativo."),
            ("N configurazioni",
             "Numero di combinazioni di parametri caricate. "
             "Piu configurazioni si testano, piu robusto e' il calcolo del PBO.",
             "Minimo: 4 file CSV | Consigliato: 8-16 o piu' | "
             "Ogni file CSV caricato = 1 configurazione",
             "Min 4 file CSV"),
            ("Partizioni S",
             f"Numero di sotto-periodi in cui viene divisa la serie storica. "
             f"Con S={s_val} si calcolano C({s_val},{s_val}//2) = {int(n_combos):,} "
             "combinazioni IS/OOS simmetriche.",
             f"S={s_val} -> {int(n_combos):,} combinazioni | "
             "Aumentare S = piu combinazioni = piu precisione "
             "ma richiede piu mesi di dati disponibili.",
             f"Attuale: S={s_val}"),
            ("Come funziona",
             "1. Si costruisce la matrice T x N (mesi x configurazioni) "
             "2. Si divide in S partizioni uguali "
             "3. Per ogni combinazione IS/OOS si identifica il parametro migliore IS "
             "e si verifica se e' sub-ottimale OOS "
             "4. PBO = frequenza in cui il miglior parametro IS non e' il migliore OOS",
             "Bailey, Borwein, Lopez de Prado, Zhu - J. Computational Finance, 2015",
             "Metodo scientifico"),
        ]

        t_leg2 = doc.add_table(rows=len(legend_l2) + 1, cols=4)
        t_leg2.style = "Table Grid"
        for j, h_txt in enumerate(["Parametro", "Cosa misura",
                                    "Come si legge", "PASS / FAIL"]):
            _bold_cell(t_leg2.cell(0, j), h_txt, bg="2E5BBA",
                       color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        for i, (param, cosa, come, esito) in enumerate(legend_l2, start=1):
            bg_r = "F8FAFC" if i % 2 == 0 else "FFFFFF"
            for j in range(4):
                _set_cell_bg(t_leg2.cell(i, j), bg_r)
            _normal_cell(t_leg2.cell(i, 0), param, bold=True)
            _normal_cell(t_leg2.cell(i, 1), cosa)
            _normal_cell(t_leg2.cell(i, 2), come)
            c3 = t_leg2.cell(i, 3)
            p3 = c3.paragraphs[0]
            for k, line in enumerate(esito.split(" | ")):
                is_pass = line.upper().startswith("PASS")
                is_fail = line.upper().startswith("FAIL")
                r3 = p3.add_run(line)
                r3.font.size = Pt(9)
                if is_pass:
                    r3.bold = True
                    r3.font.color.rgb = RGBColor.from_string("02A87A")
                elif is_fail:
                    r3.bold = True
                    r3.font.color.rgb = RGBColor.from_string("E84D4D")
                if k < len(esito.split(" | ")) - 1:
                    p3.add_run(" | ")
        doc.add_paragraph()

        # Nota conclusiva CSCV
        p_nota = doc.add_paragraph()
        rn = p_nota.add_run(
            "IMPORTANTE: La CSCV valida la robustezza del processo di selezione "
            "dei parametri (Livello 2). Non sostituisce la validazione statistica "
            "Morrone (Livello 1): entrambi i livelli devono essere superati."
        )
        rn.bold = True; rn.font.size = Pt(9)
        rn.font.color.rgb = RGBColor(0x1A, 0x2B, 0x5E)
        doc.add_paragraph()

        # ── Footer ───────────────────────────────────────────
        fp = doc.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rf = fp.add_run(
            f"Generato da Trading Strategy Validator v8  |  "
            f"{datetime.now().strftime('%d/%m/%Y %H:%M')}  |  "
            "Bailey & De Prado (2015)  -  Harvey & Liu (2014)  -  Marco Morrone"
        )
        rf.font.size = Pt(8)
        rf.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        return buf.getvalue(), None

    except Exception as e:
        import traceback
        return None, f"{e}\n{traceback.format_exc()}"



def combined_export_docx(entry):
    """
    Genera un documento Word completo con L1 (Morrone) + L2 (CSCV).
    Usato dal Tab 3 Archivio quando entrambi i livelli sono disponibili.
    Restituisce (bytes, None) oppure (None, errore).
    """
    if not DOCX_OK:
        return None, "Libreria python-docx non installata. Riavvia il .bat."
    try:
        meta  = entry.get("metadata", {})
        mor   = entry.get("morrone",  {})
        cscv  = entry.get("cscv",     None)
        m_is  = mor.get("metrics_is",  {}) or {}
        m_oos = mor.get("metrics_oos", {}) or {}
        tests = mor.get("tests",       {}) or {}
        n_p   = meta.get("n_passed", 0)
        verd  = meta.get("verdetto",
                         {5:"VALIDA",4:"QUASI VALIDA",3:"PARZIALE",
                          2:"NON IDONEA",1:"NON IDONEA",0:"NON IDONEA"
                         }.get(n_p,"NON IDONEA"))

        doc = DocxDocument()
        for sec in doc.sections:
            sec.top_margin    = Cm(2); sec.bottom_margin = Cm(2)
            sec.left_margin   = Cm(2.5); sec.right_margin  = Cm(2.5)
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)

        # ════════════════════════════════════════════════
        # INTESTAZIONE
        # ════════════════════════════════════════════════
        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = h.add_run("TRADING STRATEGY VALIDATOR v8")
        r.bold = True; r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0x1A, 0x2B, 0x5E)

        h2 = doc.add_paragraph()
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = h2.add_run("Report Completo — Livello 1 (Morrone) + Livello 2 (CSCV)")
        r2.font.size = Pt(11)
        r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        doc.add_paragraph()

        # ════════════════════════════════════════════════
        # RIEPILOGO VERDETTI
        # ════════════════════════════════════════════════
        doc.add_heading("Riepilogo Validazione", level=1)
        t_sum = doc.add_table(rows=2, cols=3)
        t_sum.style = "Table Grid"
        _bold_cell(t_sum.cell(0,0), "Parametro",  bg="1A2B5E", color="FFFFFF",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _bold_cell(t_sum.cell(0,1), "Livello 1 — Morrone", bg="1A2B5E",
                   color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        _bold_cell(t_sum.cell(0,2), "Livello 2 — CSCV", bg="1A2B5E",
                   color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)

        color_l1 = ("02A87A" if verd=="VALIDA"
                    else ("F59E0B" if "PARZIALE" in verd else "E84D4D"))
        _normal_cell(t_sum.cell(1,0), "Verdetto", bold=True)
        c_l1 = t_sum.cell(1,1)
        p_l1 = c_l1.paragraphs[0]
        p_l1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_l1 = p_l1.add_run(f"{verd} ({n_p}/5)")
        r_l1.bold = True; r_l1.font.color.rgb = RGBColor.from_string(color_l1)

        if cscv:
            pbo = float(cscv.get("pbo",0))
            color_l2 = ("02A87A" if pbo<0.05 else ("F59E0B" if pbo<0.15 else "E84D4D"))
            esito_l2 = ("ROBUSTO" if pbo<0.05 else ("ATTENZIONE" if pbo<0.15 else "OVERFITTING"))
            c_l2 = t_sum.cell(1,2)
            p_l2 = c_l2.paragraphs[0]
            p_l2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_l2 = p_l2.add_run(f"{esito_l2} (PBO={pbo:.1%})")
            r_l2.bold = True; r_l2.font.color.rgb = RGBColor.from_string(color_l2)
        else:
            _normal_cell(t_sum.cell(1,2), "Non disponibile",
                         align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

        # ════════════════════════════════════════════════
        # DATI STRATEGIA
        # ════════════════════════════════════════════════
        doc.add_heading("Dati Strategia", level=1)
        t1 = doc.add_table(rows=2, cols=4); t1.style = "Table Grid"
        for i,(h_txt,v_txt) in enumerate(zip(
            ["Strategia","Strumento","Capitale Iniziale","Data Fine IS"],
            [meta.get("strategy_name","—"), meta.get("instrument","—"),
             f"${meta.get('initial_capital',0):,}", meta.get("is_end_date","—")]
        )):
            _bold_cell(t1.cell(0,i), h_txt, bg="1A2B5E", color="FFFFFF",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _normal_cell(t1.cell(1,i), v_txt, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()
        t2 = doc.add_table(rows=2, cols=3); t2.style = "Table Grid"
        for i,(h_txt,v_txt) in enumerate(zip(
            ["N Combinazioni testate","Archiviato il","Note"],
            [f"{meta.get('n_totale',0):,}",
             meta.get("archived_at","—")[:10],
             meta.get("note","—") or "—"]
        )):
            _bold_cell(t2.cell(0,i), h_txt, bg="1A2B5E", color="FFFFFF",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _normal_cell(t2.cell(1,i), v_txt, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

        # ════════════════════════════════════════════════
        # LIVELLO 1 — TEST STATISTICI
        # ════════════════════════════════════════════════
        doc.add_heading("Livello 1 — Test Statistici (Framework Morrone)", level=1)
        if tests:
            t3 = doc.add_table(rows=len(tests)+1, cols=4); t3.style = "Table Grid"
            for j,h_txt in enumerate(["Test","Valore","Soglia","Esito"]):
                _bold_cell(t3.cell(0,j), h_txt, bg="2E5BBA", color="FFFFFF",
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            for i,(tk,tv) in enumerate(tests.items(), start=1):
                passed = tv.get("passed",False)
                _normal_cell(t3.cell(i,0), tv.get("name",tk))
                _normal_cell(t3.cell(i,1), str(tv.get("value","—")),
                             align=WD_ALIGN_PARAGRAPH.CENTER)
                _normal_cell(t3.cell(i,2), str(tv.get("soglia","—")),
                             align=WD_ALIGN_PARAGRAPH.CENTER)
                c3 = t3.cell(i,3)
                _set_cell_bg(c3, "E8F5E9" if passed else "FFEBEE")
                p3 = c3.paragraphs[0]; p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r3 = p3.add_run("PASS" if passed else "FAIL")
                r3.bold = True; r3.font.size = Pt(10)
                r3.font.color.rgb = RGBColor.from_string("02A87A" if passed else "E84D4D")
        doc.add_paragraph()

        # Metriche IS/OOS
        doc.add_heading("Metriche Performance", level=2)
        has_oos = bool(m_oos)
        metr_rows = [
            ("Trade",        str(int(m_is.get("n_trades",0))),
             str(int(m_oos.get("n_trades",0))) if has_oos else "—"),
            ("Sharpe Ratio", f"{float(m_is.get('sharpe',0)):.3f}",
             f"{float(m_oos.get('sharpe',0)):.3f}" if has_oos else "—"),
            ("CAGR",         f"{float(m_is.get('cagr',0)):.2%}",
             f"{float(m_oos.get('cagr',0)):.2%}" if has_oos else "—"),
            ("Max Drawdown", f"{float(m_is.get('max_dd',0)):.2%}",
             f"{float(m_oos.get('max_dd',0)):.2%}" if has_oos else "—"),
            ("Profit Factor",f"{float(m_is.get('profit_factor',0)):.2f}",
             f"{float(m_oos.get('profit_factor',0)):.2f}" if has_oos else "—"),
            ("Win Rate",     f"{float(m_is.get('win_rate',0)):.1%}",
             f"{float(m_oos.get('win_rate',0)):.1%}" if has_oos else "—"),
            ("Avg Trade ($)",f"${float(m_is.get('avg_trade',0)):.0f}",
             f"${float(m_oos.get('avg_trade',0)):.0f}" if has_oos else "—"),
        ]
        t4 = doc.add_table(rows=len(metr_rows)+1, cols=3); t4.style = "Table Grid"
        for j,h_txt in enumerate(["Metrica","In-Sample","Out-of-Sample"]):
            _bold_cell(t4.cell(0,j), h_txt, bg="2E5BBA", color="FFFFFF",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for i,(label,is_v,oos_v) in enumerate(metr_rows, start=1):
            bg_r = "F8FAFC" if i%2==0 else "FFFFFF"
            for j in range(3): _set_cell_bg(t4.cell(i,j), bg_r)
            _normal_cell(t4.cell(i,0), label, bold=True)
            _normal_cell(t4.cell(i,1), is_v,  align=WD_ALIGN_PARAGRAPH.CENTER)
            _normal_cell(t4.cell(i,2), oos_v, align=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()

        # Legenda L1
        doc.add_heading("Legenda Test Statistici", level=2)
        legend_l1 = [
            ("T-Statistic",
             "Il profitto medio e' statisticamente diverso da zero?",
             "t > 1.96 (alpha=5%)",
             "PASS se t > 1.96 | FAIL se t <= 1.96"),
            ("E[MaxSR] / EMSR",
             "Lo Sharpe e' davvero buono o solo il migliore per caso tra N tentativi?",
             "SR osservato > E[MaxSR] atteso da N strategie casuali",
             "PASS se SR > E[MaxSR] | FAIL altrimenti"),
            ("MinBTL",
             "Anni minimi di backtest necessari per N ottimizzazioni effettuate.",
             "Anni IS >= MinBTL calcolato",
             "PASS se Anni IS >= MinBTL | FAIL altrimenti"),
            ("Bonferroni",
             "Correzione per test multipli: con N ottimizzazioni la soglia sale.",
             "t-stat > t_critico(alpha/N)",
             "PASS se t > soglia | FAIL altrimenti"),
            ("DSR",
             "Deflated Sharpe Ratio: corregge per N, non-normalita' e lunghezza.",
             "DSR > soglia (default 0.95)",
             "PASS se DSR > 0.95 | FAIL altrimenti"),
        ]
        t_l1 = doc.add_table(rows=len(legend_l1)+1, cols=4); t_l1.style = "Table Grid"
        for j,h_txt in enumerate(["Test","Cosa misura","Come si legge","PASS / FAIL"]):
            _bold_cell(t_l1.cell(0,j), h_txt, bg="1A2B5E", color="FFFFFF",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for i,(test,cosa,come,esito) in enumerate(legend_l1, start=1):
            bg_r = "F8FAFC" if i%2==0 else "FFFFFF"
            for j in range(4): _set_cell_bg(t_l1.cell(i,j), bg_r)
            _normal_cell(t_l1.cell(i,0), test, bold=True)
            _normal_cell(t_l1.cell(i,1), cosa)
            _normal_cell(t_l1.cell(i,2), come)
            c3 = t_l1.cell(i,3); p3 = c3.paragraphs[0]
            for k,line in enumerate(esito.split(" | ")):
                is_pass = line.upper().startswith("PASS")
                is_fail = line.upper().startswith("FAIL")
                r3 = p3.add_run(line); r3.font.size = Pt(9)
                if is_pass: r3.bold=True; r3.font.color.rgb=RGBColor.from_string("02A87A")
                elif is_fail: r3.bold=True; r3.font.color.rgb=RGBColor.from_string("E84D4D")
                if k < len(esito.split(" | "))-1: p3.add_run(" | ")
        doc.add_paragraph()
        p_v = doc.add_paragraph()
        rv = p_v.add_run("Verdetto: 5/5=VALIDA | 4/5=QUASI VALIDA | 3/5=PARZIALE | 2/5 o meno=NON IDONEA")
        rv.bold=True; rv.font.size=Pt(9); rv.font.color.rgb=RGBColor(0x1A,0x2B,0x5E)
        doc.add_paragraph()

        # ════════════════════════════════════════════════
        # LIVELLO 2 — CSCV
        # ════════════════════════════════════════════════
        if cscv:
            doc.add_heading("Livello 2 — CSCV (Combinatorially Symmetric Cross-Validation)", level=1)
            pr = doc.add_paragraph()
            rr = pr.add_run("Bailey, Borwein, Lopez de Prado, Zhu — J. Computational Finance, 2015")
            rr.italic=True; rr.font.size=Pt(9)
            rr.font.color.rgb=RGBColor(0x94,0xA3,0xB8)

            pbo  = float(cscv.get("pbo",0))
            pl   = float(cscv.get("prob_loss",0))
            esito_c = ("ROBUSTO (PBO < 5%)" if pbo<0.05
                       else ("ATTENZIONE (5-15%)" if pbo<0.15 else "OVERFITTING (> 15%)"))
            color_c = ("02A87A" if pbo<0.05 else ("F59E0B" if pbo<0.15 else "E84D4D"))
            pc = doc.add_paragraph(); pc.alignment=WD_ALIGN_PARAGRAPH.CENTER
            rc = pc.add_run(f"PBO = {pbo:.1%}  —  {esito_c}")
            rc.bold=True; rc.font.size=Pt(14)
            rc.font.color.rgb=RGBColor.from_string(color_c)

            cscv_rows = [
                ("PBO",                    f"{pbo:.1%}",                        "< 5%"),
                ("Prob. perdita OOS",       f"{pl:.1%}",                         "< 30%"),
                ("N config. testate",       str(cscv.get("N_configs","—")),      "Min 4"),
                ("N combinazioni IS/OOS",   f"{int(cscv.get('n_combos',0)):,}", "—"),
                ("Mesi (T)",                str(cscv.get("T_months","—")),       "—"),
                ("Partizioni (S)",          str(cscv.get("S","—")),              "16"),
            ]
            t5 = doc.add_table(rows=len(cscv_rows)+1, cols=3); t5.style="Table Grid"
            for j,h_txt in enumerate(["Parametro","Valore","Soglia"]):
                _bold_cell(t5.cell(0,j), h_txt, bg="2E5BBA", color="FFFFFF",
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            for i,(label,val,soglia) in enumerate(cscv_rows, start=1):
                bg_r = "F8FAFC" if i%2==0 else "FFFFFF"
                for j in range(3): _set_cell_bg(t5.cell(i,j), bg_r)
                _normal_cell(t5.cell(i,0), label, bold=True)
                _normal_cell(t5.cell(i,1), val,    align=WD_ALIGN_PARAGRAPH.CENTER)
                _normal_cell(t5.cell(i,2), soglia, align=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_paragraph()

            # Legenda L2
            doc.add_heading("Legenda Parametri CSCV", level=2)
            legend_l2 = [
                ("PBO",
                 "Prob. che il parametro migliore IS sia peggiore del benchmark OOS.",
                 "< 5%=robusto | 5-15%=attenzione | > 15%=overfitting",
                 "PASS se PBO < 5% | FAIL se PBO >= 5%"),
                ("Prob. perdita OOS",
                 "Frequenza con cui la strategia selezionata IS perde nei periodi OOS.",
                 "Piu bassa e' meglio. Riferimento: < 30%",
                 "Nessuna soglia fissa"),
                ("Beta degradazione",
                 "Pendenza SR-IS vs SR-OOS. Beta < 0 = SR alto IS predice SR basso OOS.",
                 "Beta > 0: buono | Beta = 0: neutro | Beta < 0: overfitting classico",
                 "PASS se Beta > 0 | FAIL se Beta < 0"),
                ("N configurazioni",
                 "File CSV caricati. Piu configurazioni = calcolo PBO piu robusto.",
                 "Min 4 | Consigliato 8-16+",
                 "Min 4 file CSV"),
            ]
            t6 = doc.add_table(rows=len(legend_l2)+1, cols=4); t6.style="Table Grid"
            for j,h_txt in enumerate(["Parametro","Cosa misura","Come si legge","PASS / FAIL"]):
                _bold_cell(t6.cell(0,j), h_txt, bg="1A2B5E", color="FFFFFF",
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            for i,(param,cosa,come,esito) in enumerate(legend_l2, start=1):
                bg_r = "F8FAFC" if i%2==0 else "FFFFFF"
                for j in range(4): _set_cell_bg(t6.cell(i,j), bg_r)
                _normal_cell(t6.cell(i,0), param, bold=True)
                _normal_cell(t6.cell(i,1), cosa)
                _normal_cell(t6.cell(i,2), come)
                c3 = t6.cell(i,3); p3 = c3.paragraphs[0]
                for k,line in enumerate(esito.split(" | ")):
                    is_pass=line.upper().startswith("PASS")
                    is_fail=line.upper().startswith("FAIL")
                    r3=p3.add_run(line); r3.font.size=Pt(9)
                    if is_pass: r3.bold=True; r3.font.color.rgb=RGBColor.from_string("02A87A")
                    elif is_fail: r3.bold=True; r3.font.color.rgb=RGBColor.from_string("E84D4D")
                    if k < len(esito.split(" | "))-1: p3.add_run(" | ")
            doc.add_paragraph()
            p_nota = doc.add_paragraph()
            rn = p_nota.add_run(
                "IMPORTANTE: La CSCV valida il processo di selezione parametri (L2). "
                "Non sostituisce la validazione statistica Morrone (L1): entrambi devono essere superati.")
            rn.bold=True; rn.font.size=Pt(9); rn.font.color.rgb=RGBColor(0x1A,0x2B,0x5E)
            doc.add_paragraph()

        # ════════════════════════════════════════════════
        # FOOTER
        # ════════════════════════════════════════════════
        fp = doc.add_paragraph(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rf = fp.add_run(
            f"Report Completo L1+L2 — Trading Strategy Validator v8  |  "
            f"{datetime.now().strftime('%d/%m/%Y %H:%M')}  |  "
            "Bailey & De Prado (2015)  -  Harvey & Liu (2014)  -  Marco Morrone")
        rf.font.size=Pt(8); rf.font.color.rgb=RGBColor(0x94,0xA3,0xB8)

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        return buf.getvalue(), None

    except Exception as e:
        import traceback
        return None, f"{e}\n{traceback.format_exc()}"

# ══════════════════════════════════════════════════════════════
# PAGINA
# ══════════════════════════════════════════════════════════════
# Forza il limite upload a 2 GB indipendentemente da come viene avviato
from streamlit import config as _st_config
try:
    _st_config.set_option("server.maxUploadSize", 2048)
except Exception:
    pass

st.set_page_config(
    page_title="Trading Strategy Validator v8",
    page_icon="📊", layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #1A2B5E 0%, #2E5BBA 100%);
        padding: 2rem; border-radius: 12px; margin-bottom: 2rem; text-align: center;
    }
    .main-header h1 { color: white; font-size: 2.2rem; margin: 0; }
    .main-header p  { color: #C8D8F8; font-size: 1rem; margin: 0.5rem 0 0 0; }
    .metric-card {
        background: white; border-radius: 10px; padding: 1.2rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        border-left: 4px solid #2E5BBA; margin-bottom: 1rem;
        color: #1A2B5E !important;
    }
    .metric-card * { color: #1A2B5E !important; }
    .metric-card.green { border-left-color: #02A87A; }
    .metric-card.red   { border-left-color: #E84D4D; }
    .verdict-pass {
        background: linear-gradient(135deg, #02A87A, #028060);
        color: white !important; padding: 2rem; border-radius: 12px;
        text-align: center; font-size: 1.8rem; font-weight: bold; margin: 1rem 0;
    }
    .verdict-fail {
        background: linear-gradient(135deg, #E84D4D, #C03030);
        color: white !important; padding: 2rem; border-radius: 12px;
        text-align: center; font-size: 1.8rem; font-weight: bold; margin: 1rem 0;
    }
    .verdict-partial {
        background: linear-gradient(135deg, #F59E0B, #D97706);
        color: white !important; padding: 2rem; border-radius: 12px;
        text-align: center; font-size: 1.8rem; font-weight: bold; margin: 1rem 0;
    }
    .info-box {
        background: #EEF3FC; border: 1px solid #2E5BBA;
        border-radius: 8px; padding: 1rem; margin: 0.5rem 0; font-size: 0.9rem;
        color: #1A2B5E !important;
    }
    .info-box * { color: #1A2B5E !important; }
    .info-warn {
        background: #FFF8E1; border: 1px solid #F59E0B;
        border-radius: 8px; padding: 1rem; margin: 0.5rem 0; font-size: 0.9rem;
        color: #92400E !important;
    }
    .info-warn * { color: #92400E !important; }
    .section-header {
        font-size: 1.3rem; font-weight: bold; color: #1A2B5E !important;
        border-bottom: 2px solid #2E5BBA; padding-bottom: 0.5rem;
        margin: 1.5rem 0 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="main-header">
    <h1>📊 Trading Strategy Validator v8</h1>
    <p>Framework Morrone · CSCV Bailey &amp; De Prado · Archivio locale strategie</p>
</div>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## ⚙️ Configurazione")
    st.markdown("---")
    st.markdown("### 📋 Strategia")
    strategy_name   = st.text_input("Nome strategia", value="", placeholder="Es: Crabel ORB NQ 15min")
    initial_capital = st.number_input("Capitale iniziale ($)",
                                       min_value=1000, max_value=10_000_000,
                                       value=100_000, step=1000)
    instrument = st.text_input("Strumento", value="NQ Futures")

    st.markdown("### 📅 Fine In-Sample")
    is_end_date = st.date_input(
        "Data fine In-Sample",
        value=pd.to_datetime("2023-12-31"),
        help="Trade fino a questa data = IS. Dopo = OOS."
    )

    st.markdown("### 🔢 Combinazioni testate (N)")
    n_mode = st.radio("Modalità:",
                      ["Totale diretto", "Per fasi (calcolo automatico)"])
    if n_mode == "Totale diretto":
        n_totale = st.number_input("N totale", min_value=1,
                                    max_value=10_000_000, value=367)
    else:
        n_fasi = st.number_input("Fasi di ottimizzazione", 1, 10, 3)
        totali_fasi = []
        dettaglio   = []
        for fi in range(int(n_fasi)):
            fn  = fi + 1
            nm  = st.text_input(f"Nome fase {fn}", value=f"Fase {fn}",
                                  key=f"nf{fi}", label_visibility="collapsed")
            np_ = st.number_input(f"Parametri in {nm}", 1, 8, 2, key=f"np{fi}")
            cols_p = st.columns(min(int(np_), 4))
            prod, pv = 1, []
            for pi in range(int(np_)):
                with cols_p[pi % 4]:
                    v = st.number_input(f"P{pi+1}", 1, 100_000, 15,
                                         key=f"f{fi}p{pi}")
                    prod *= v; pv.append(v)
            totali_fasi.append(prod)
            dettaglio.append((nm, pv, prod))
            st.caption(f"→ {' × '.join(str(x) for x in pv)} = **{prod:,}**")
            st.markdown("---")
        n_totale = sum(totali_fasi)
        st.success(f"**N = {' + '.join(str(d[2]) for d in dettaglio)} = {n_totale:,}**")

    with st.expander("⚙️ Parametri avanzati"):
        alpha         = st.slider("Alpha", 0.01, 0.10, 0.05, 0.01)
        dsr_threshold = st.slider("Soglia DSR", 0.80, 0.99, 0.95, 0.01)
        trading_days  = st.number_input("Giorni trading/anno", 200, 365, 252)
        s_cscv        = st.number_input("Partizioni CSCV (S)", 4, 24, 16, 2,
                                          help="S=16 → 12.870 combo IS/OOS")

    st.markdown("---")
    st.markdown("<small>📚 Bailey & De Prado (2014, 2015)<br>"
                "Harvey & Liu (2014) · Marco Morrone</small>",
                unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════
# FUNZIONI — CARICAMENTO
# ══════════════════════════════════════════════════════════════
def load_csv(f, cap):
    """
    CSV formato standard:
    separatore=;  decimale=,  date=DD/MM/YYYY HH:MM:SS
    Colonne obbligatorie: exitTime, profit
    """
    try:
        content = f.read().decode("utf-8", errors="replace")
        df = pd.read_csv(io.StringIO(content), sep=";", decimal=",")
        miss = {"profit","exitTime"} - set(df.columns)
        if miss:
            return None, f"Colonne mancanti: {miss}. Presenti: {list(df.columns)}"

        df["exit_datetime"]  = pd.to_datetime(df["exitTime"],
                                               format="%d/%m/%Y %H:%M:%S",
                                               errors="coerce")
        df["entry_datetime"] = pd.to_datetime(
            df.get("entryTime", df["exitTime"]),
            format="%d/%m/%Y %H:%M:%S", errors="coerce"
        )
        df["profit"] = pd.to_numeric(
            df["profit"].astype(str).str.replace(",", "."), errors="coerce"
        )
        df = df.dropna(subset=["profit","exit_datetime"]).reset_index(drop=True)
        df = df.sort_values("exit_datetime").reset_index(drop=True)
        if len(df) == 0:
            return None, "Nessun trade valido."

        df["cum_profit"]  = df["profit"].cumsum()
        df["equity"]      = cap + df["cum_profit"]
        df["equity_prev"] = df["equity"].shift(1).fillna(cap)
        df["return_pct"]  = df["profit"] / df["equity_prev"]
        df["trade_num"]   = range(1, len(df)+1)
        return df, None
    except Exception as e:
        import traceback
        return None, f"{e}\n{traceback.format_exc()}"


def load_xls(f, cap):
    """XLS MultiCharts — foglio List of Trades."""
    try:
        xl    = pd.ExcelFile(f)
        sheet = "List of Trades" if "List of Trades" in xl.sheet_names \
                else xl.sheet_names[0]
        raw = pd.read_excel(f, sheet_name=sheet, header=None)
        hr  = 0
        for i in range(min(10, len(raw))):
            vs = [str(v).lower() for v in raw.iloc[i].values]
            if any("trade" in v for v in vs) and any("profit" in v for v in vs):
                hr = i; break
        df       = pd.read_excel(f, sheet_name=sheet, header=hr)
        col_list = list(df.columns)

        def fc(names):
            for nm in names:
                for c in col_list:
                    if str(c).strip().lower() == nm.lower():
                        return c
            return None

        ct  = fc(["Type"])
        cd  = fc(["Date"])
        cp  = fc(["Profit ($)"])
        ccp = fc(["Cum. Profit ($)"])
        miss = [n for n,c in [("Type",ct),("Date",cd),
                               ("Profit ($)",cp),("Cum. Profit ($)",ccp)] if c is None]
        if miss:
            return None, f"Colonne non trovate: {miss}\nPresenti: {col_list}"

        work = pd.DataFrame({
            "type":       df[ct].astype(str).str.strip(),
            "date":       pd.to_datetime(df[cd], errors="coerce"),
            "profit":     pd.to_numeric(df[cp], errors="coerce"),
            "cum_profit": pd.to_numeric(df[ccp], errors="coerce"),
        })
        em = work["type"].isin(["EntryLong","EntryShort"])
        xm = work["type"].isin(["ExitLong","ExitShort"])
        de = work[em].reset_index(drop=True)
        dx = work[xm].reset_index(drop=True)
        if len(de) == 0:
            return None, f"Nessuna riga Entry. Tipi trovati: {work['type'].unique()}"
        n  = min(len(de), len(dx))
        de = de.iloc[:n].reset_index(drop=True)
        dx = dx.iloc[:n].reset_index(drop=True)
        trades = pd.DataFrame({
            "trade_num":      range(1, n+1),
            "entry_datetime": de["date"].values,
            "exit_datetime":  dx["date"].values,
            "profit":         de["profit"].values,
            "cum_profit":     de["cum_profit"].values,
        }).dropna(subset=["profit"]).sort_values("exit_datetime").reset_index(drop=True)
        trades["equity"]      = cap + trades["cum_profit"]
        trades["equity_prev"] = trades["equity"].shift(1).fillna(cap)
        trades["return_pct"]  = trades["profit"] / trades["equity_prev"]
        return trades, None
    except Exception as e:
        import traceback
        return None, f"{e}\n{traceback.format_exc()}"


def load_csv_monthly(f, label=""):
    """
    Legge P&L mensile da tre formati:

    1. CSV con Indice (nuovo formato):
       Header: Indice,Data,Profit
       sep=,  dec=.  Data=seriale Excel
       → Mantiene indici primo/ultimo trade per mese

    2. TXT PowerLanguage (senza header):
       Due colonne spazio-separate: serial_date profit
       → Stessa conversione data seriale Excel

    3. CSV standard (formato originale):
       Header con exitTime e profit
       sep=;  dec=,  exitTime=DD/MM/YYYY HH:MM:SS
    """
    try:
        fname   = f.name.lower()
        content = f.read().decode("utf-8", errors="replace")
        lines   = [l.strip() for l in content.strip().splitlines() if l.strip()]

        # ── Formato 1: CSV con Indice,Data,Profit (o exitTime,profit) ──────────
        if lines and lines[0].lower().startswith("indice"):
            import io as _io
            df = pd.read_csv(_io.StringIO(content), sep=",", decimal=".")
            # Normalizza nomi colonne (case-insensitive)
            df.columns = [c.strip().lower() for c in df.columns]
            # Accetta sia 'data' che 'exittime' come colonna data
            date_col = next((c for c in df.columns if c in ("data", "exittime")), None)
            if date_col is None or "profit" not in df.columns:
                return None, (f"'{label}': colonne 'Data' e 'Profit' non trovate. "
                              f"Presenti: {df.columns.tolist()}")
            df["date"] = (pd.Timestamp("1899-12-30")
                          + pd.to_timedelta(df[date_col], unit="D"))
            df["month"] = df["date"].dt.to_period("M").astype(str)

            # Aggrega per mese mantenendo gli indici
            has_idx = "indice" in df.columns
            if has_idx:
                monthly = df.groupby("month").agg(
                    profit    = ("profit", "sum"),
                    n_trades  = ("indice", "count"),
                    idx_first = ("indice", "first"),
                    idx_last  = ("indice", "last"),
                ).rename(columns={"profit": label})
                # Per il CSCV serve solo la serie P&L mensile;
                # gli indici vengono conservati come colonne extra
                monthly.columns.name = None
                result = monthly[label].copy()
                result._idx_info = monthly[["n_trades","idx_first","idx_last"]]
            else:
                result = df.groupby("month")["profit"].sum().rename(label)
            if len(result) == 0:
                return None, "EMPTY"   # file con solo header: skip silenzioso
            # Conta trade totali: usa _idx_info se disponibile, altrimenti len(df)
            if has_idx:
                result._n_trades = int(monthly["n_trades"].sum())
            else:
                result._n_trades = len(df)
            return result, None

        # ── Formato 2: TXT PowerLanguage (senza header) ─────
        if fname.endswith(".txt") or (
            len(lines) > 0 and len(lines[0].split()) == 2
            and not lines[0][0].isalpha()
        ):
            records = []
            for line in lines:
                parts = line.split()
                if len(parts) >= 2:
                    try:
                        serial = float(parts[0])
                        profit = float(parts[1])
                        dt = (pd.Timestamp("1899-12-30")
                              + pd.to_timedelta(serial, unit="D"))
                        records.append({"date": dt, "profit": profit})
                    except Exception:
                        continue
            if not records:
                return None, "EMPTY"   # nessuna riga dati: skip silenzioso
            df = pd.DataFrame(records)
            df["month"] = df["date"].dt.to_period("M").astype(str)
            result = df.groupby("month")["profit"].sum().rename(label)
            if len(result) == 0:
                return None, "EMPTY"
            result._n_trades = len(records)
            return result, None

        # ── Formato 3: CSV standard exitTime;profit ──────────
        import io as _io
        df = pd.read_csv(_io.StringIO(content), sep=";", decimal=",")
        if "profit" not in df.columns or "exitTime" not in df.columns:
            return None, (f"'{label}': formato non riconosciuto.\n"
                          f"Colonne trovate: {df.columns.tolist()}\n"
                          f"Formati supportati:\n"
                          f"  1. CSV con header Indice,Data,Profit\n"
                          f"  2. TXT con due colonne: serial_date profit\n"
                          f"  3. CSV con exitTime e profit (sep=;)")
        df["exit_dt"] = pd.to_datetime(df["exitTime"],
                                        format="%d/%m/%Y %H:%M:%S",
                                        errors="coerce")
        df["profit"]  = pd.to_numeric(
            df["profit"].astype(str).str.replace(",", "."), errors="coerce"
        )
        df = df.dropna(subset=["exit_dt", "profit"])
        if len(df) == 0:
            return None, "EMPTY"   # header presente ma 0 trade validi: skip silenzioso
        df["month"] = df["exit_dt"].dt.to_period("M").astype(str)
        result = df.groupby("month")["profit"].sum().rename(label)
        result._n_trades = len(df)
        return result, None

    except Exception as e:
        return None, f"Errore '{label}': {e}"


# ══════════════════════════════════════════════════════════════
# FUNZIONI — CALCOLO MORRONE
# ══════════════════════════════════════════════════════════════
def compute_metrics(df, cap, label=""):
    ret  = df["return_pct"].dropna()
    pnl  = df["profit"].dropna()
    n    = len(ret)
    if n < 5: return None
    years   = (df["exit_datetime"].max()-df["exit_datetime"].min()).days/365.25
    tpy     = n/years if years > 0 else 0
    std_ret = ret.std(ddof=1)
    sharpe  = ret.mean()/std_ret*np.sqrt(tpy) if std_ret > 0 else 0
    feq     = df["equity"].iloc[-1]
    cagr    = (feq/cap)**(1/years)-1 if years > 0 else 0
    rm      = df["equity"].cummax()
    max_dd  = ((df["equity"]-rm)/rm).min()
    gp = pnl[pnl>0].sum(); gl = abs(pnl[pnl<0].sum())
    return {
        "label": label, "n_trades": n, "years": years, "tpy": tpy,
        "sharpe": sharpe, "cagr": cagr, "max_dd": max_dd,
        "profit_factor": gp/gl if gl>0 else np.inf,
        "win_rate": (pnl>0).mean(), "avg_trade": pnl.mean(),
        "final_equity": feq,
        "skewness": stats.skew(ret), "kurtosis": stats.kurtosis(ret),
        "returns": ret, "profits": pnl,
    }


def calc_emsr(n_strat, n_obs):
    ge = 0.5772156649
    N, T = max(n_strat,2), max(n_obs,2)
    return ((1-ge)*stats.norm.ppf(1-1/N)+ge*stats.norm.ppf(1-1/(N*np.e)))/np.sqrt(T)


def run_all_tests(m, n_strat, alpha, dsr_threshold):
    results = {}
    t       = m["sharpe"]*np.sqrt(m["years"])
    p       = 2*(1-stats.norm.cdf(abs(t)))
    ts      = stats.norm.ppf(1-alpha/2)
    results["ttest"] = {
        "name":"T-Statistic","value":round(t,3),"soglia":round(ts,3),
        "extra":f"P-value = {p:.4f}","passed":t>ts,"formula":"T = Sharpe × √anni"
    }
    sr0t = calc_emsr(n_strat, m["n_trades"])
    sr0a = sr0t*np.sqrt(m["tpy"])
    results["emsr"] = {
        "name":"Sharpe IS > E[MaxSR]","value":round(m["sharpe"],3),
        "soglia":round(sr0a,3),"extra":f"SR casuale con N={n_strat}: {sr0a:.3f}",
        "passed":m["sharpe"]>sr0a,"formula":"E[MaxSR] = f(N, T)"
    }
    try:
        minbtl = brentq(lambda y: calc_emsr(n_strat,y*m["tpy"])*np.sqrt(m["tpy"])-1.0,
                        0.01, 300)
    except Exception:
        minbtl = float("nan")
    results["minbtl"] = {
        "name":"Anni IS >= MinBTL","value":round(m["years"],1),
        "soglia":round(minbtl,1),"extra":f"Minimi per N={n_strat}: {minbtl:.1f}",
        "passed":m["years"]>=minbtl if not np.isnan(minbtl) else False,
        "formula":"MinBTL = f(N)"
    }
    ab = alpha/n_strat; sb = stats.norm.ppf(1-ab/2)
    results["bonferroni"] = {
        "name":"T-stat > Bonferroni","value":round(t,3),"soglia":round(sb,2),
        "extra":f"α aggiustato = {alpha:.2f}/{n_strat} = {ab:.5f}",
        "passed":t>sb,"formula":"α_bonf = α / N"
    }
    srt   = m["sharpe"]/np.sqrt(m["tpy"]) if m["tpy"]>0 else 0
    T     = m["n_trades"]
    sigma = np.sqrt(max(1e-10,(1-m["skewness"]*srt+
                               ((m["kurtosis"]+2)/4)*srt**2)/max(T-1,1)))
    z   = (srt-sr0t)/sigma; dsr = stats.norm.cdf(z)
    results["dsr"] = {
        "name":"DSR > soglia","value":round(dsr,3),"soglia":dsr_threshold,
        "extra":f"Corregge per N, Skew={m['skewness']:.2f}, Kurt={m['kurtosis']:.2f}",
        "passed":dsr>dsr_threshold,"formula":"DSR = Φ(z)"
    }
    return results, minbtl, sr0a


def make_charts(df, df_is, df_oos, m_is):
    fig = plt.figure(figsize=(14,10))
    gs  = gridspec.GridSpec(2,2,figure=fig,hspace=0.4,wspace=0.35)
    NAVY,BLUE,RED,GRAY = "#1A2B5E","#2E5BBA","#E84D4D","#64748B"
    ax1 = fig.add_subplot(gs[0,:])
    ax1.plot(df_is["exit_datetime"],df_is["equity"],color=BLUE,lw=1.8,
             label=f"In-Sample ({len(df_is)} trade)")
    if df_oos is not None and len(df_oos)>0:
        ax1.plot(df_oos["exit_datetime"],df_oos["equity"],color=RED,lw=1.8,
                 label=f"Out-of-Sample ({len(df_oos)} trade)")
        ax1.axvline(df_is["exit_datetime"].iloc[-1],color=GRAY,ls="--",alpha=0.6,lw=1.2)
    ax1.set_title("Equity Curve",fontsize=13,fontweight="bold",color=NAVY,pad=10)
    ax1.set_ylabel("Equity ($)",color=GRAY)
    ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda x,_:f"${x:,.0f}"))
    ax1.legend(fontsize=9); ax1.grid(True,alpha=0.25)
    ax1.spines[["top","right"]].set_visible(False)
    ax2 = fig.add_subplot(gs[1,0])
    rm  = df["equity"].cummax()
    dd  = (df["equity"]-rm)/rm*100
    ax2.fill_between(df["exit_datetime"],dd,0,color=RED,alpha=0.35)
    ax2.set_title("Drawdown Storico",fontsize=11,fontweight="bold",color=NAVY)
    ax2.set_ylabel("Drawdown (%)",color=GRAY)
    ax2.grid(True,alpha=0.25); ax2.spines[["top","right"]].set_visible(False)
    ax3 = fig.add_subplot(gs[1,1])
    ret = m_is["returns"]
    ax3.hist(ret,bins=30,density=True,color=BLUE,alpha=0.55,label="Rendimenti IS")
    x = np.linspace(ret.min(),ret.max(),200)
    ax3.plot(x,stats.norm.pdf(x,ret.mean(),ret.std()),color=RED,lw=2,label="Normale")
    ax3.axvline(0,color=GRAY,ls="--",alpha=0.5)
    ax3.set_title(f"Distribuzione IS\n(Skew={m_is['skewness']:.2f}, "
                  f"Kurt={m_is['kurtosis']:.2f})",
                  fontsize=11,fontweight="bold",color=NAVY)
    ax3.legend(fontsize=9); ax3.grid(True,alpha=0.25)
    ax3.spines[["top","right"]].set_visible(False)
    fig.patch.set_facecolor("#F8FAFC")
    for ax in [ax1,ax2,ax3]: ax.set_facecolor("#FFFFFF")
    return fig


# ══════════════════════════════════════════════════════════════
# FUNZIONI — CSCV
# ══════════════════════════════════════════════════════════════
def sharpe_monthly(s):
    if len(s)<3: return 0.0
    mu,std = s.mean(),s.std(ddof=1)
    return (mu/std)*np.sqrt(12) if std>0 else 0.0


def run_cscv(matrix_df, S=16):
    T,N = matrix_df.shape
    if T<S: return None,f"Periodi ({T}) < S ({S}). Riduci S."
    if N<4: return None,f"Servono almeno 4 configurazioni (trovate {N})."
    sz     = T//S
    M      = matrix_df.iloc[:sz*S].values
    slices = [M[i*sz:(i+1)*sz,:] for i in range(S)]
    idx_S  = list(range(S))
    logits,sr_is_l,sr_oos_l = [],[],[]
    for is_idx in combinations(idx_S, S//2):
        oos_idx = [i for i in idx_S if i not in is_idx]
        IS  = np.vstack([slices[i] for i in is_idx])
        OOS = np.vstack([slices[i] for i in oos_idx])
        sr_is = np.array([sharpe_monthly(IS[:,n]) for n in range(N)])
        ns    = int(np.argmax(sr_is))
        sr_oos_all = np.array([sharpe_monthly(OOS[:,n]) for n in range(N)])
        sr_oos_ns  = sharpe_monthly(OOS[:,ns])
        omega = np.clip(np.sum(sr_oos_all<=sr_oos_ns)/(N+1),1e-6,1-1e-6)
        logits.append(np.log(omega/(1-omega)))
        sr_is_l.append(sr_is[ns]); sr_oos_l.append(sr_oos_ns)
    logits    = np.array(logits)
    sr_is_arr = np.array(sr_is_l)
    sr_oos_arr= np.array(sr_oos_l)
    return {
        "pbo":float(np.mean(logits<0)), "logits":logits,
        "sr_is_list":sr_is_arr, "sr_oos_list":sr_oos_arr,
        "prob_loss":float(np.mean(sr_oos_arr<0)),
        "N_configs":N,"T_months":sz*S,"S":S,"n_combos":len(logits),
    }, None


def compute_neff(matrix_df, threshold=0.95):
    """
    Calcola N_eff tramite PCA sulla matrice mensile P&L (mesi × configurazioni).

    Algoritmo:
    1. Normalizza ogni colonna (media zero, varianza unitaria)
    2. Calcola la matrice di covarianza N×N
    3. Estrae gli autovalori (decomposizione spettrale)
    4. Ordina in modo decrescente e calcola la varianza spiegata cumulata
    5. Neff = numero minimo di componenti per raggiungere 'threshold' di varianza

    Restituisce: (neff, eigenvalues_sorted, cumulative_explained_ratio)
    """
    M = matrix_df.fillna(0).values.astype(float)
    T, N = M.shape
    if N < 2:
        return 1, np.array([1.0]), np.array([1.0])

    # Normalizzazione: sottrai media e dividi per std per ogni configurazione
    std_col = M.std(axis=0, ddof=1)
    std_col[std_col == 0] = 1.0          # evita divisione per zero
    M_norm = (M - M.mean(axis=0)) / std_col

    # Matrice di covarianza N×N
    cov = np.cov(M_norm.T)               # shape (N, N)
    if cov.ndim == 0:                    # edge case N=1
        return 1, np.array([float(cov)]), np.array([1.0])

    # Autovalori (reali perché cov è simmetrica)
    eigenvalues = np.linalg.eigvalsh(cov)
    eigenvalues = np.maximum(eigenvalues, 0)   # clip errori numerici
    eigenvalues = np.sort(eigenvalues)[::-1]   # ordine decrescente

    total_var = eigenvalues.sum()
    if total_var == 0:
        return 1, eigenvalues, np.ones(len(eigenvalues))

    cumulative = np.cumsum(eigenvalues) / total_var
    neff = int(np.searchsorted(cumulative, threshold) + 1)
    neff = max(1, min(neff, N))          # bounded [1, N]
    return neff, eigenvalues, cumulative


def make_cscv_charts(cr):
    NAVY,BLUE,RED,GRAY,GREEN = "#1A2B5E","#2E5BBA","#E84D4D","#64748B","#02A87A"
    fig,axes = plt.subplots(1,2,figsize=(14,6))
    fig.patch.set_facecolor("#F8FAFC")
    pbo,logits = cr["pbo"],cr["logits"]
    sr_is,sr_oos,pl = cr["sr_is_list"],cr["sr_oos_list"],cr["prob_loss"]
    ax1 = axes[0]; ax1.set_facecolor("white")
    nb  = min(50,max(20,len(logits)//100))
    ax1.hist(logits,bins=nb,density=True,color=BLUE,alpha=0.65)
    mu_l,sd_l = logits.mean(),logits.std()
    xr = np.linspace(logits.min(),logits.max(),300)
    ax1.plot(xr,stats.norm.pdf(xr,mu_l,sd_l),color=RED,lw=2,ls="--",label="Normale")
    ax1.axvline(0,color=RED,lw=2,alpha=0.8)
    ax1.fill_between(xr[xr<0],stats.norm.pdf(xr[xr<0],mu_l,sd_l),alpha=0.15,color=RED)
    ax1.set_title(f"Distribuzione Logit\nPBO={pbo:.1%}  |  Prob.perdita OOS={pl:.1%}",
                  fontsize=12,fontweight="bold",color=NAVY)
    ax1.set_xlabel("λ = ln[ω/(1-ω)]",color=GRAY); ax1.set_ylabel("Frequenza",color=GRAY)
    ax1.text(0.05,0.95,f"PBO={pbo:.1%}\n({'✓ <5%' if pbo<0.05 else '✗ ≥5%'})",
             transform=ax1.transAxes,fontsize=13,fontweight="bold",
             color=GREEN if pbo<0.05 else RED,va="top",
             bbox=dict(boxstyle="round,pad=0.4",facecolor="white",
                       edgecolor=GREEN if pbo<0.05 else RED,alpha=0.9))
    ax1.legend(fontsize=9); ax1.grid(True,alpha=0.25)
    ax1.spines[["top","right"]].set_visible(False)
    ax2 = axes[1]; ax2.set_facecolor("white")
    ax2.scatter(sr_is,sr_oos,alpha=0.12,s=8,color=BLUE)
    if len(sr_is)>10:
        try:
            sl,ic,rv,_,_ = stats.linregress(sr_is,sr_oos)
            xf = np.linspace(sr_is.min(),sr_is.max(),100)
            ax2.plot(xf,sl*xf+ic,color=RED,lw=2,label=f"β={sl:.2f}, R²={rv**2:.2f}")
        except Exception: pass
    ax2.axhline(0,color=GRAY,ls="--",alpha=0.5,lw=1)
    ax2.axvline(0,color=GRAY,ls="--",alpha=0.5,lw=1)
    ax2.set_title(f"Performance Degradation IS→OOS\nProb(SR_OOS<0)={pl:.1%}",
                  fontsize=12,fontweight="bold",color=NAVY)
    ax2.set_xlabel("Sharpe IS",color=GRAY); ax2.set_ylabel("Sharpe OOS",color=GRAY)
    ax2.legend(fontsize=9); ax2.grid(True,alpha=0.25)
    ax2.spines[["top","right"]].set_visible(False)
    plt.tight_layout(pad=2.0); return fig


# ══════════════════════════════════════════════════════════════
# TAB
# ══════════════════════════════════════════════════════════════
tab_m, tab_c, tab_a = st.tabs([
    "📊 Livello 1 — Framework Morrone",
    "🔄 Livello 2 — CSCV (Bailey & De Prado 2015)",
    "📦 Archivio Strategie",
])


# ══════════════════════════════════════════════════════════════
# TAB 1 — MORRONE
# ══════════════════════════════════════════════════════════════
with tab_m:
    cu, ci = st.columns([2,1])
    with cu:
        st.markdown('<div class="section-header">📁 Carica Trade List</div>',
                    unsafe_allow_html=True)
        fmt = st.radio("Formato:", ["XLS MultiCharts","CSV (formato standard)"],
                       horizontal=True, key="fmt_m")
        if "CSV" in fmt:
            uploaded = st.file_uploader("File CSV (sep=; decimal=,)",
                                         type=["csv"], key="up_m_csv")
        else:
            uploaded = st.file_uploader("File XLS MultiCharts",
                                         type=["xlsx","xls"], key="up_m_xls")
    with ci:
        st.markdown('<div class="section-header">ℹ️ Formato CSV</div>',
                    unsafe_allow_html=True)
        st.markdown("""
**Colonne richieste:**
- `exitTime` — `DD/MM/YYYY HH:MM:SS`
- `profit` — P&L per trade
- Separatore: `;` · Decimale: `,`

**Colonne opzionali:**
- `entryTime`, `type` (Long/Short), `tradeN`
        """)

    if uploaded is not None:
        if "CSV" in fmt:
            trades, err = load_csv(uploaded, initial_capital)
        else:
            trades, err = load_xls(uploaded, initial_capital)

        if err:
            st.error(f"❌ {err}")
        else:
            is_end = pd.Timestamp(is_end_date)
            df_is  = trades[trades["exit_datetime"]<=is_end].copy().reset_index(drop=True)
            df_oos = trades[trades["exit_datetime"]> is_end].copy().reset_index(drop=True)

            if len(df_is)==0:
                st.error("Nessun trade nel periodo IS. Controlla la data fine IS.")
                st.stop()

            df_is["cum_p"]        = df_is["profit"].cumsum()
            df_is["equity"]       = initial_capital + df_is["cum_p"]
            df_is["equity_prev"]  = df_is["equity"].shift(1).fillna(initial_capital)
            df_is["return_pct"]   = df_is["profit"] / df_is["equity_prev"]

            if len(df_oos)>0:
                leq = df_is["equity"].iloc[-1]
                df_oos["cum_p"]       = df_oos["profit"].cumsum()
                df_oos["equity"]      = leq + df_oos["cum_p"]
                df_oos["equity_prev"] = df_oos["equity"].shift(1).fillna(leq)
                df_oos["return_pct"]  = df_oos["profit"] / df_oos["equity_prev"]

            m_is  = compute_metrics(df_is, initial_capital, "IS")
            m_oos = compute_metrics(df_oos, leq, "OOS") \
                    if len(df_oos)>5 else None
            if not m_is:
                st.error("Dati IS insufficienti."); st.stop()

            test_results, minbtl, sr0a = run_all_tests(
                m_is, n_totale, alpha, dsr_threshold
            )
            n_passed = sum(1 for t in test_results.values() if t["passed"])

            ti_gr, ti_test, ti_met, ti_rep = st.tabs(
                ["📈 Grafici","🧪 Test","📊 Metriche","📄 Report & Archivio"]
            )

            with ti_gr:
                fig_m = make_charts(trades, df_is, df_oos, m_is)
                st.pyplot(fig_m, use_container_width=True)
                plt.close(fig_m)

            with ti_test:
                vmap2 = {5:"VALIDA",4:"QUASI VALIDA",3:"PARZIALE",
                         2:"NON IDONEA",1:"NON IDONEA",0:"NON IDONEA"}
                if n_passed==5:
                    st.markdown('<div class="verdict-pass">✅ VALIDA — 5/5</div>',
                                unsafe_allow_html=True)
                elif n_passed>=3:
                    st.markdown(f'<div class="verdict-partial">⚠️ PARZIALE — {n_passed}/5</div>',
                                unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="verdict-fail">✗ NON IDONEA — {n_passed}/5</div>',
                                unsafe_allow_html=True)
                st.markdown("")
                for tk, t in test_results.items():
                    icon = "✅" if t["passed"] else "❌"
                    cls  = "green" if t["passed"] else "red"
                    c1,c2 = st.columns([3,1])
                    with c1:
                        st.markdown(
                            f'<div class="metric-card {cls}">'
                            f'<b>{icon} {t["name"]}</b><br>'
                            f'<small>{t["formula"]}</small><br>'
                            f'<small style="color:#64748B">{t["extra"]}</small>'
                            f'</div>', unsafe_allow_html=True)
                    with c2:
                        st.metric("Valore / Soglia", str(t["value"]),
                                  delta=f"soglia: {t['soglia']}")

            with ti_met:
                rows_m = [
                    ("Trade",         str(m_is["n_trades"]),   str(m_oos["n_trades"]) if m_oos else "—"),
                    ("Sharpe",        f"{m_is['sharpe']:.3f}", f"{m_oos['sharpe']:.3f}" if m_oos else "—"),
                    ("CAGR",          f"{m_is['cagr']:.2%}",   f"{m_oos['cagr']:.2%}" if m_oos else "—"),
                    ("Max Drawdown",  f"{m_is['max_dd']:.2%}", f"{m_oos['max_dd']:.2%}" if m_oos else "—"),
                    ("Profit Factor", f"{m_is['profit_factor']:.2f}",
                                      f"{m_oos['profit_factor']:.2f}" if m_oos else "—"),
                    ("Win Rate",      f"{m_is['win_rate']:.1%}",f"{m_oos['win_rate']:.1%}" if m_oos else "—"),
                    ("Avg Trade ($)", f"${m_is['avg_trade']:.0f}",
                                      f"${m_oos['avg_trade']:.0f}" if m_oos else "—"),
                    ("Skewness",      f"{m_is['skewness']:.3f}", "—"),
                    ("Kurtosis",      f"{m_is['kurtosis']:.3f}", "—"),
                ]
                st.dataframe(pd.DataFrame(rows_m,
                    columns=["Metrica","In-Sample","Out-of-Sample"]),
                    use_container_width=True, hide_index=True)

            with ti_rep:
                vmap = {5:"VALIDA",4:"QUASI VALIDA",3:"PARZIALE",
                        2:"NON IDONEA",1:"NON IDONEA",0:"NON IDONEA"}
                lines = [
                    "="*55,
                    "  TRADING STRATEGY VALIDATOR v8",
                    f"  Strategia:  {strategy_name}",
                    f"  Strumento:  {instrument}",
                    f"  Capitale:   ${initial_capital:,.0f}",
                    f"  IS fino al: {is_end_date}",
                    "="*55,"",
                    "METRICHE IN-SAMPLE","-"*55,
                    f"  Trade:      {m_is['n_trades']}",
                    f"  Sharpe:     {m_is['sharpe']:.3f}",
                    f"  CAGR:       {m_is['cagr']:.2%}",
                    f"  Max DD:     {m_is['max_dd']:.2%}",
                    f"  PF:         {m_is['profit_factor']:.2f}",
                    f"  Win Rate:   {m_is['win_rate']:.1%}",
                    f"  Avg Trade:  ${m_is['avg_trade']:.0f}",
                    f"  Skewness:   {m_is['skewness']:.3f}",
                    f"  Kurtosis:   {m_is['kurtosis']:.3f}","",
                    f"PARAMETRI (N={n_totale}, alpha={alpha:.0%})","-"*55,
                    f"  MinBTL:     {minbtl:.2f} anni",
                    f"  E[MaxSR]:   {sr0a:.3f}","",
                    "TEST STATISTICI","-"*55,
                ]
                for t in test_results.values():
                    lines.append(f"  [{'OK' if t['passed'] else 'NO'}] "
                                 f"{t['name']:<28} {t['value']} / {t['soglia']}")
                if m_oos:
                    deg = (m_oos["sharpe"]/m_is["sharpe"]-1)*100 \
                          if m_is["sharpe"]!=0 else 0
                    lines += ["","METRICHE OUT-OF-SAMPLE","-"*55,
                               f"  Trade:      {m_oos['n_trades']}",
                               f"  Sharpe:     {m_oos['sharpe']:.3f}",
                               f"  CAGR:       {m_oos['cagr']:.2%}",
                               f"  PF:         {m_oos['profit_factor']:.2f}",
                               f"  Avg Trade:  ${m_oos['avg_trade']:.0f}",
                               f"  Degr. SR:   {deg:+.1f}%"]
                lines += ["","="*55,
                          f"  VERDETTO: {vmap.get(n_passed,'NON IDONEA')} ({n_passed}/5)",
                          "="*55]
                report_text = "\n".join(lines)
                st.code(report_text, language="")

                cd1,cd2 = st.columns(2)
                with cd1:
                    st.download_button("💾 Scarica Report (.txt)",
                        data=report_text,
                        file_name=f"report_{strategy_name.replace(' ','_')}.txt",
                        mime="text/plain")
                with cd2:
                    fig2 = make_charts(trades, df_is, df_oos, m_is)
                    buf  = io.BytesIO()
                    fig2.savefig(buf,format="png",dpi=150,
                                 bbox_inches="tight",facecolor="#F8FAFC")
                    plt.close(fig2); buf.seek(0)
                    st.download_button("🖼️ Scarica Grafici (.png)",
                        data=buf,
                        file_name=f"grafici_{strategy_name.replace(' ','_')}.png",
                        mime="image/png")

                st.markdown("---")
                # ── Auto-salvataggio archivio ─────────────────
                _save_key = f"saved_{strategy_name}_{n_passed}_{len(trades)}"
                if strategy_name and st.session_state.get("_l1_save_key") != _save_key:
                    try:
                        note_arc = ""
                        eid = arc_save_morrone(
                            strategy_name, instrument, initial_capital,
                            is_end_date, n_totale, alpha, dsr_threshold,
                            trades, m_is, m_oos, test_results, n_passed,
                            note=note_arc
                        )
                        st.session_state["last_arc_id"]       = eid
                        st.session_state["last_arc_strategy"]  = strategy_name
                        st.session_state["_l1_save_key"]       = _save_key
                        st.success(f"✅ **Livello 1 salvato automaticamente** in archivio — {strategy_name}")
                    except Exception as e:
                        st.error(f"Errore salvataggio automatico: {e}")
                elif strategy_name:
                    st.success(f"✅ **Livello 1 già in archivio** — {strategy_name}  "
                               f"→ vai al Tab 2 per eseguire la CSCV")
                else:
                    st.warning("⚠️ Inserisci il **Nome strategia** nella sidebar per salvare in archivio.")

    else:
        st.markdown("---")
        c1,c2,c3 = st.columns(3)
        for col,num,tit,desc in [
            (c1,"1️⃣","Configura sidebar","Capitale, data IS/OOS e N combinazioni"),
            (c2,"2️⃣","Esporta da MultiCharts",
             "CSV (sep=; dec=,) oppure XLS → List of Trades"),
            (c3,"3️⃣","Carica e analizza","Verdetto completo con tutti i 5 test"),
        ]:
            col.markdown(f"""
            <div style="background:white;border-radius:10px;padding:1.5rem;
                        box-shadow:0 2px 8px rgba(0,0,0,0.08);text-align:center;
                        color:#1A2B5E;">
                <div style="font-size:2rem;">{num}</div>
                <div style="font-weight:bold;color:#1A2B5E;margin:0.5rem 0;">{tit}</div>
                <div style="color:#64748B;font-size:0.9rem;">{desc}</div>
            </div>""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════
# TAB 2 — CSCV
# ══════════════════════════════════════════════════════════════
with tab_c:
    CSCV_KEY = "cscv_last_result"  # definita qui per essere disponibile in tutto il tab
    st.markdown("""
    <div class="info-box">
    <b>🔄 CSCV — Combinatorially Symmetric Cross-Validation</b><br>
    Bailey, Borwein, López de Prado, Zhu — J. Computational Finance, 2015<br><br>
    Carica un <b>file CSV per ogni configurazione</b> testata durante l'ottimizzazione.
    Il programma aggrega il P&L per mese, costruisce la matrice T×N e calcola
    il <b>PBO</b> su migliaia di combinazioni IS/OOS simmetriche.
    Soglia: <b>PBO &lt; 5%</b>
    </div>
    """, unsafe_allow_html=True)

    with st.expander("📋 Formati accettati e naming file"):
        st.markdown("""
**Formato TXT PowerLanguage (quello che hai):**
- Due colonne separate da spazio, senza header
- Colonna 1: data in formato **seriale Excel** (es. 42010.31)
- Colonna 2: profit del trade in dollari (es. 1200.00 o -437.50)
```
42010.31  1200.00
42017.31  2687.50
42031.31 -1437.50
```

**Formato CSV standard (alternativo):**
- Separatore `;` — decimale `,`
- Colonne: `exitTime` (DD/MM/YYYY HH:MM:SS) e `profit`

**Un file per ogni configurazione testata.** Il programma ordina i file
alfabeticamente → usa naming che rispetti l'ordine logico:

```
cscv_fase_001_PtnNeuLY_001.txt    ← valore 1
cscv_fase_001_PtnNeuLY_002.txt    ← valore 2
...
cscv_fase_004_MyProfit_01000.txt
...
```
        """)

    st.markdown('<div class="section-header">📁 Carica i CSV (uno per configurazione)</div>',
                unsafe_allow_html=True)

    # Gestione tasto svuotamento: cambiare la key resetta il file_uploader
    if "cscv_uploader_key" not in st.session_state:
        st.session_state["cscv_uploader_key"] = 0

    col_up, col_clr = st.columns([5, 1])
    with col_clr:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🗑️ Svuota file", key="btn_clear_cscv",
                     help="Rimuove tutti i file CSV caricati"):
            st.session_state["cscv_uploader_key"] += 1
            st.rerun()

    with col_up:
        uploaded_cscv = st.file_uploader(
            "Trascina qui tutti i file delle configurazioni testate (.csv / .txt / .zip)",
            type=["csv", "txt", "zip"], accept_multiple_files=True,
            key=f"up_cscv_{st.session_state['cscv_uploader_key']}"
        )

    if uploaded_cscv:
        # ── Estrazione ZIP ────────────────────────────────────
        import zipfile, io as _io2
        expanded_files = []
        for uf in uploaded_cscv:
            if uf.name.lower().endswith(".zip"):
                try:
                    with zipfile.ZipFile(_io2.BytesIO(uf.read())) as zf:
                        csv_names = sorted([
                            n for n in zf.namelist()
                            if not n.startswith("__MACOSX")
                            and not n.startswith(".")
                            and n.lower().endswith((".csv", ".txt"))
                            and "/" not in n.rstrip("/").replace("\\", "/").rsplit("/", 1)[0]
                               or True  # accetta anche sottocartelle
                        ])
                        # Filtra solo file (non directory) e non __MACOSX
                        csv_names = [
                            n for n in zf.namelist()
                            if not n.startswith("__MACOSX")
                            and not n.endswith("/")
                            and n.lower().endswith((".csv", ".txt"))
                        ]
                        for name in sorted(csv_names):
                            data = zf.read(name)
                            fake = _io2.BytesIO(data)
                            fake.name = name.split("/")[-1]  # solo nome file
                            expanded_files.append(fake)
                    st.info(f"📦 ZIP **{uf.name}**: estratti {len(expanded_files)} file CSV/TXT")
                except Exception as e:
                    st.error(f"Errore lettura ZIP {uf.name}: {e}")
            else:
                expanded_files.append(uf)

        n_files = len(expanded_files)
        st.info(f"📂 {n_files} file caricati — elaborati in ordine alfabetico per nome.")
        files_sorted = sorted(expanded_files, key=lambda f: f.name)

        # ── Controlli soglia filtro ───────────────────────────
        with st.expander("⚙️ Filtri file — soglia minima trade", expanded=False):
            st.markdown("""
            Esclude automaticamente i file con pochi trade prima del calcolo CSCV.
            I file con solo header (0 trade) vengono sempre saltati senza errore.
            Puoi impostare una soglia **assoluta** (numero di trade) oppure **relativa**
            (% rispetto al file con più trade). Viene applicata la soglia più restrittiva
            tra le due attive.
            """)
            fc1, fc2 = st.columns(2)
            with fc1:
                thresh_abs = st.number_input(
                    "Trade minimi assoluti per file",
                    min_value=0, value=0, step=1,
                    help="0 = nessun filtro assoluto. Es: 10 esclude file con < 10 trade."
                )
            with fc2:
                thresh_pct = st.slider(
                    "% minima rispetto al file con più trade",
                    min_value=0, max_value=100, value=0, step=1,
                    help="0 = nessun filtro percentuale. Es: 20 esclude file con meno del 20% dei trade del file più ricco."
                )

        with st.spinner("Lettura P&L mensile..."):
            # Fase 1: carica tutto
            raw_series, raw_labels, errors, empty_files = [], [], [], []
            idx_info_list = []

            for f in files_sorted:
                label = f.name.replace(".csv", "").replace(".txt", "")
                s, err = load_csv_monthly(f, label)
                if err == "EMPTY":
                    empty_files.append(f.name)
                elif err:
                    errors.append(f"**{f.name}**: {err}")
                else:
                    raw_series.append(s)
                    raw_labels.append(label)
                    if hasattr(s, "_idx_info"):
                        idx_info_list.append((label, s._idx_info))

            # Fase 2: applica soglie
            n_trades_list = [getattr(s, "_n_trades", 0) for s in raw_series]
            max_trades = max(n_trades_list) if n_trades_list else 1
            pct_threshold_abs = (thresh_pct / 100.0) * max_trades

            series_list, labels, filtered_files = [], [], []
            for s, lbl, nt in zip(raw_series, raw_labels, n_trades_list):
                effective_min = max(thresh_abs, pct_threshold_abs)
                if effective_min > 0 and nt < effective_min:
                    filtered_files.append((lbl, nt))
                else:
                    series_list.append(s)
                    labels.append(lbl)

        # Riepilogo caricamento
        if empty_files:
            with st.expander(f"⚪ {len(empty_files)} file vuoti ignorati (solo header)"):
                for fn in empty_files:
                    st.markdown(f"- `{fn}`")

        if errors:
            with st.expander(f"⚠️ {len(errors)} file con errori di parsing"):
                for e in errors:
                    st.markdown(f"- {e}")

        if filtered_files:
            with st.expander(f"🔽 {len(filtered_files)} file esclusi dalla soglia "
                             f"(min {int(max(thresh_abs, pct_threshold_abs))} trade)"):
                for lbl, nt in filtered_files:
                    pct_of_max = nt / max_trades * 100 if max_trades > 0 else 0
                    st.markdown(f"- `{lbl}` — **{nt}** trade ({pct_of_max:.1f}% del massimo)")

        if n_trades_list and (thresh_abs > 0 or thresh_pct > 0):
            st.caption(
                f"File massimo trade: **{max_trades}** | "
                f"Soglia attiva: ≥ **{int(max(thresh_abs, pct_threshold_abs))}** trade | "
                f"File ammessi: **{len(series_list)}** / {len(raw_series)}"
            )

        # ── Diagnostica caricamento ───────────────────────────
        _can_run = True

        if len(series_list) < 2:
            _can_run = False
            n_empty   = len(empty_files)
            n_errors  = len(errors)
            n_filtered = len(filtered_files)
            st.error(
                f"**File validi insufficienti per il CSCV** (trovati {len(series_list)}, "
                f"servono almeno 2).\n\n"
                f"Riepilogo caricamento:\n"
                f"- 📂 File totali caricati: **{len(files_sorted)}**\n"
                f"- ✅ File validi con dati: **{len(series_list)}**\n"
                f"- ⚪ File vuoti (solo header): **{n_empty}**\n"
                f"- ⚠️ File con errori parsing: **{n_errors}**\n"
                f"- 🔽 File esclusi dalla soglia: **{n_filtered}**\n\n"
                f"**Possibili cause:** i CSV esportati da MultiCharts potrebbero avere "
                f"0 trade per quella configurazione di parametri, oppure il formato "
                f"non è riconosciuto. Prova ad abbassare la soglia filtro a 0."
            )

        if _can_run:
            matrix_df = pd.DataFrame(series_list).T.sort_index()
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Configurazioni (N)", len(series_list))
            c2.metric("Mesi totali",        len(matrix_df))
            c3.metric("Mesi comuni",        int(matrix_df.dropna().shape[0]))
            pct_na = matrix_df.isna().sum().sum() / matrix_df.size
            c4.metric("Dati mancanti",      f"{pct_na:.1%}")

            na_method = st.selectbox("Gestione mesi mancanti:",
                ["Solo mesi comuni a tutte (consigliato)", "Tutti i mesi, mancanti → 0"])
            matrix_clean = matrix_df.dropna() if "comuni" in na_method \
                            else matrix_df.fillna(0)

            S_val = int(s_cscv)
            if len(matrix_clean) < S_val:
                _can_run = False
                st.error(
                    f"**Mesi disponibili insufficienti:** {len(matrix_clean)} mesi comuni "
                    f"< S={S_val} partizioni richieste.\n\n"
                    f"Soluzioni:\n"
                    f"- Riduci S nella sidebar (prova S=8)\n"
                    f"- Seleziona **'Tutti i mesi, mancanti → 0'** nel menu sopra\n"
                    f"- Carica configurazioni con più storia comune"
                )
            else:
                st.markdown('<div class="section-header">Preview matrice P&L mensile</div>',
                            unsafe_allow_html=True)
                st.dataframe(matrix_clean.tail(12).style.format("{:.0f}"),
                             use_container_width=True)
                st.caption(f"Ultimi 12 mesi. Matrice: {len(matrix_clean)} mesi × "
                           f"{len(labels)} configurazioni.")

                # ── Calcolo e visualizzazione Neff ───────────────────
                st.markdown('<div class="section-header">🔬 N Effettivo (Neff) — PCA sulla matrice mensile</div>',
                            unsafe_allow_html=True)
                neff_val, neff_eig, neff_cum = None, None, None
                try:
                    neff_val, neff_eig, neff_cum = compute_neff(matrix_clean, threshold=0.95)
                    n_nominal = len(series_list)
                    pct_reduction = (1 - neff_val / n_nominal) * 100 if n_nominal > 0 else 0

                    st.markdown(f"""
                    <div class="info-box">
                    <b>N nominale = {n_nominal}</b> configurazioni &nbsp;|&nbsp;
                    <b>Neff = {neff_val}</b> componenti indipendenti &nbsp;|&nbsp;
                    Riduzione: <b>{pct_reduction:.0f}%</b>
                    </div>
                    """, unsafe_allow_html=True)

                    nc1, nc2, nc3 = st.columns(3)
                    nc1.metric("N nominale", n_nominal)
                    nc2.metric("Neff (95%)", neff_val,
                               delta=f"−{n_nominal - neff_val} correlate",
                               delta_color="off")
                    _cum_val = (f"{neff_cum[neff_val-1]:.1%}"
                                if neff_val >= 1 and neff_val <= len(neff_cum) else "—")
                    nc3.metric("Var. spiegata a Neff", _cum_val)

                    with st.expander("📊 Dettaglio PCA"):
                        n_show = min(30, len(neff_eig))
                        _tot_var = neff_eig.sum()
                        df_pca = pd.DataFrame({
                            "Componente":        list(range(1, n_show + 1)),
                            "Autovalore":        neff_eig[:n_show].round(3),
                            "Var. spiegata (%)": (neff_eig[:n_show] / _tot_var * 100).round(2)
                                                 if _tot_var > 0 else [0]*n_show,
                            "Cumulata (%)":      (neff_cum[:n_show] * 100).round(2),
                        })
                        df_pca["In Neff"] = df_pca["Componente"] <= neff_val
                        st.dataframe(df_pca, use_container_width=True, hide_index=True)

                    st.info(
                        f"I test Bonferroni, MinBTL e DSR verranno ricalcolati con "
                        f"**Neff={neff_val}** dopo l'esecuzione CSCV."
                    )
                except Exception as _neff_err:
                    neff_val = None
                    st.warning(f"⚠️ Neff non calcolabile: {_neff_err}. "
                               f"Il CSCV procederà con N nominale.")

                # ── Indici trade ──────────────────────────────────────
                if idx_info_list:
                    with st.expander(f"📋 Indici trade per mese ({len(idx_info_list)} configurazioni)"):
                        label0, idx0 = idx_info_list[0]
                        st.markdown(f"**Esempio — `{label0}`:**")
                        st.dataframe(idx0.rename(columns={
                            "n_trades":"Trade","idx_first":"Primo","idx_last":"Ultimo"}),
                            use_container_width=True)
                        rows_idx = []
                        for lbl, idf in idx_info_list:
                            tmp = idf.copy().reset_index()
                            tmp.insert(0, "Configurazione", lbl)
                            rows_idx.append(tmp)
                        if rows_idx:
                            df_all_idx = pd.concat(rows_idx, ignore_index=True)
                            st.download_button(
                                "💾 Scarica indici trade per mese (.csv)",
                                data=df_all_idx.to_csv(index=False, sep=";", decimal=","),
                                file_name="cscv_indici_trade_mese.csv", mime="text/csv")

        # ── Sezione pulsante — SEMPRE VISIBILE ───────────────
        st.markdown('<div class="section-header">Esecuzione CSCV</div>',
                    unsafe_allow_html=True)

        if not _can_run:
            st.warning("⚠️ Risolvi i problemi segnalati sopra per abilitare il calcolo.")
            btn_col1, btn_col2 = st.columns([4, 1])
            with btn_col1:
                st.button("▶️ Esegui CSCV", type="primary",
                          use_container_width=True, disabled=True)
            with btn_col2:
                if st.button("🗑️ Pulisci", use_container_width=True,
                             help="Cancella risultati in memoria"):
                    st.session_state.pop(CSCV_KEY, None)
                    st.rerun()
        else:
            n_combos = comb(S_val, S_val//2)
            st.markdown(f"""
            <div class="info-box">
            <b>Parametri:</b> S={S_val} → <b>{n_combos:,} combinazioni IS/OOS</b> &nbsp;|&nbsp;
            N={len(series_list)} config &nbsp;|&nbsp;
            T={len(matrix_clean)} mesi &nbsp;|&nbsp;
            Partizione ~ {len(matrix_clean)//S_val} mesi
            </div>
            """, unsafe_allow_html=True)

            # Collegamento archivio — automatico da L1 o per nome
            _manifest_c = arc_manifest()
            arc_link_id = st.session_state.get("last_arc_id")
            _arc_strategy = st.session_state.get("last_arc_strategy", "")
            if arc_link_id:
                st.markdown(
                    f'<div class="info-box">&#x1F517; <b>Collegamento automatico</b> alla strategia L1: <b>{_arc_strategy}</b><br>'
                    f'I risultati CSCV verranno salvati automaticamente nella stessa cartella.</div>',
                    unsafe_allow_html=True)
            else:
                _arc_opts = ["— Seleziona strategia L1 —"] + \
                            [f"{m['strategy_name']}" for m in _manifest_c if not m.get("has_cscv")]
                _arc_sel = st.selectbox(
                    "Collega a strategia L1 in archivio:",
                    _arc_opts,
                    help="Seleziona la strategia Morrone a cui collegare questi risultati CSCV."
                )
                if _arc_sel != "— Seleziona strategia L1 —":
                    _match = next((m for m in _manifest_c if m["strategy_name"]==_arc_sel), None)
                    if _match:
                        arc_link_id = _match["id"]
                        _arc_strategy = _arc_sel

            btn_col1, btn_col2 = st.columns([4, 1])
            with btn_col1:
                run_cscv_btn = st.button("▶️ Esegui CSCV", type="primary",
                                         use_container_width=True)
            with btn_col2:
                if st.button("🗑️ Pulisci", use_container_width=True,
                             help="Cancella risultati in memoria per rieseguire"):
                    st.session_state.pop(CSCV_KEY, None)
                    st.rerun()

            # ── Calcolo ──────────────────────────────────────────
            if run_cscv_btn:
                with st.spinner(f"Calcolo {n_combos:,} combinazioni IS/OOS..."):
                    cscv_result, err = run_cscv(matrix_clean, S=S_val)
                if err:
                    st.error(f"Errore CSCV: {err}")
                else:
                    sr_is  = cscv_result["sr_is_list"]
                    sr_oos = cscv_result["sr_oos_list"]
                    try:
                        sl,ic,rv,_,_ = stats.linregress(sr_is, sr_oos)
                        beta_str, r2_str = f"{sl:.2f}", f"{rv**2:.2f}"
                    except Exception:
                        beta_str, r2_str = "N/A", "N/A"
                    pbo = cscv_result["pbo"]; pl = cscv_result["prob_loss"]
                    interp_lines = [
                        f"Configurazioni (N):  {len(series_list)}",
                        f"Mesi per partizione: {len(matrix_clean)//S_val}",
                        f"Combinazioni:        {cscv_result['n_combos']:,}", "",
                        f"PBO:                 {pbo:.1%}  ({'PASS' if pbo<0.05 else 'FAIL'})",
                        f"Prob.perdita OOS:    {pl:.1%}",
                        f"Beta degradazione:   {beta_str}",
                        f"R²:                  {r2_str}", "",
                        "GUIDA LETTURA:",
                        "  PBO < 5%   → selezione non overfittata",
                        "  PBO 5-15%  → attenzione, verificare con più configurazioni",
                        "  PBO > 15%  → overfitting significativo nel processo di selezione",
                        "  Beta < 0   → SR alto IS predice SR basso OOS (classico overfitting)",
                    ]
                    fig_c2 = make_cscv_charts(cscv_result)
                    buf_c  = io.BytesIO()
                    fig_c2.savefig(buf_c, format="png", dpi=150,
                                   bbox_inches="tight", facecolor="#F8FAFC")
                    plt.close(fig_c2); buf_c.seek(0)
                    st.session_state[CSCV_KEY] = {
                        "result":       cscv_result,
                        "matrix":       matrix_clean,
                        "n_series":     len(series_list),
                        "arc_link_id":  arc_link_id,
                        "interp_lines": interp_lines,
                        "img_bytes":    buf_c.getvalue(),
                        "neff":         neff_val if neff_val is not None else None,
                        "neff_cum":     neff_cum if neff_cum is not None else None,
                        "neff_eig":     neff_eig if neff_eig is not None else None,
                        "n_nominal":    len(series_list),
                    }
    

    else:
        st.markdown("---")
        c1,c2,c3 = st.columns(3)
        for col,num,tit,desc in [
            (c1,"1️⃣","Esporta CSV per config",
             "Un file per ogni configurazione: exitTime, profit"),
            (c2,"2️⃣","Naming ordinato",
             "fase1_ATR_03.csv, fase1_ATR_04.csv ... Il nome determina l'ordine"),
            (c3,"3️⃣","Carica tutti insieme",
             "Trascina tutti i CSV qui. Il PBO viene calcolato automaticamente"),
        ]:
            col.markdown(f"""
            <div style="background:white;border-radius:10px;padding:1.5rem;
                        box-shadow:0 2px 8px rgba(0,0,0,0.08);text-align:center;
                        color:#1A2B5E;">
                <div style="font-size:2rem;">{num}</div>
                <div style="font-weight:bold;color:#1A2B5E;margin:0.5rem 0;">{tit}</div>
                <div style="color:#64748B;font-size:0.9rem;">{desc}</div>
            </div>""", unsafe_allow_html=True)




    # ── Mostra risultati da session_state ─────────────────
    if CSCV_KEY in st.session_state:
        saved        = st.session_state[CSCV_KEY]
        cscv_result  = saved["result"]
        matrix_saved = saved["matrix"]
        arc_id_saved = saved.get("arc_link_id")
        interp_lines = saved["interp_lines"]
        img_bytes    = saved["img_bytes"]
        neff_saved   = saved.get("neff")
        n_nom_saved  = saved.get("n_nominal")

        st.info("💾 Risultati in memoria — premi **🗑️ Pulisci** per rieseguire "
                "o **▶️ Esegui CSCV** per ricalcolare con nuovi file.")

        pbo = cscv_result["pbo"]; pl = cscv_result["prob_loss"]
        if pbo < 0.05:
            st.markdown(f'<div class="verdict-pass">✅ PBO={pbo:.1%} — ROBUSTO (&lt;5%)</div>',
                        unsafe_allow_html=True)
        elif pbo < 0.15:
            st.markdown(f'<div class="verdict-partial">⚠️ PBO={pbo:.1%} — ATTENZIONE (5-15%)</div>',
                        unsafe_allow_html=True)
        else:
            st.markdown(f'<div class="verdict-fail">✗ PBO={pbo:.1%} — OVERFITTING (&gt;15%)</div>',
                        unsafe_allow_html=True)

        st.markdown("")
        cc1,cc2,cc3,cc4 = st.columns(4)
        cc1.metric("PBO", f"{pbo:.1%}",
                   delta="obiettivo <5%",
                   delta_color="inverse" if pbo >= 0.05 else "normal")
        cc2.metric("Prob.perdita OOS", f"{pl:.1%}")
        cc3.metric("Combinazioni", f"{cscv_result['n_combos']:,}")
        cc4.metric("Configurazioni", cscv_result["N_configs"])

        fig_cscv = make_cscv_charts(cscv_result)
        st.pyplot(fig_cscv, use_container_width=True)
        plt.close(fig_cscv)

        # ── Ricalcolo L1 con Neff ────────────────────────────
        if neff_saved is not None:
            st.markdown(
                '<div class="section-header">🔬 Ricalcolo Test L1 con Neff — Valutazione Definitiva</div>',
                unsafe_allow_html=True)
            st.markdown(f"""
            <div class="info-box">
            I test Bonferroni, MinBTL e DSR vengono ricalcolati usando <b>Neff = {neff_saved}</b>
            invece di N nominale = {n_nom_saved}. Questo è il giudizio
            <b>metodologicamente corretto</b>: le configurazioni correlate non costituiscono
            test indipendenti.
            </div>
            """, unsafe_allow_html=True)

            # Recupera metriche L1 dall'archivio se disponibili
            _l1_entry = None
            if arc_id_saved:
                _l1_entry, _ = arc_load(arc_id_saved)
            m_is_arc = (_l1_entry.get("morrone", {}).get("metrics_is") or {}) \
                       if _l1_entry else {}
            tests_arc = (_l1_entry.get("morrone", {}).get("tests") or {}) \
                        if _l1_entry else {}

            # Prova a ricalcolare i 3 test dipendenti da N usando Neff
            if m_is_arc and tests_arc:
                t_val     = float(tests_arc.get("ttest", {}).get("value", 0) or 0)
                years_val = float(m_is_arc.get("years", 0) or 0)
                tpy_val   = float(m_is_arc.get("tpy", 0) or 0)
                sr_val    = float(m_is_arc.get("sharpe", 0) or 0)
                sk_val    = float(m_is_arc.get("skewness", 0) or 0)
                ku_val    = float(m_is_arc.get("kurtosis", 0) or 0)
                n_tr_val  = int(m_is_arc.get("n_trades", 1) or 1)
                alpha_val = float(_l1_entry["metadata"].get("alpha", 0.05) or 0.05)
                dsr_thr   = float(_l1_entry["metadata"].get("dsr_threshold", 0.95) or 0.95)

                # --- Bonferroni con Neff ---
                ab_neff   = alpha_val / neff_saved
                sb_neff   = stats.norm.ppf(1 - ab_neff / 2)
                bon_pass_nom = bool(tests_arc.get("bonferroni", {}).get("passed", False))
                bon_pass_neff = t_val > sb_neff

                # --- MinBTL con Neff ---
                try:
                    minbtl_neff = brentq(
                        lambda y: calc_emsr(neff_saved, y * tpy_val) * np.sqrt(tpy_val) - 1.0,
                        0.01, 300
                    )
                    minbtl_pass_neff = years_val >= minbtl_neff
                except Exception:
                    minbtl_neff = float("nan")
                    minbtl_pass_neff = False
                minbtl_pass_nom = bool(tests_arc.get("minbtl", {}).get("passed", False))

                # --- DSR con Neff (sr0t ricalcolato con Neff) ---
                sr0t_neff = calc_emsr(neff_saved, n_tr_val)
                srt_val   = sr_val / np.sqrt(tpy_val) if tpy_val > 0 else 0
                sigma_dsr = np.sqrt(max(1e-10, (1 - sk_val * srt_val +
                                                 ((ku_val + 2) / 4) * srt_val ** 2) / max(n_tr_val - 1, 1)))
                z_neff    = (srt_val - sr0t_neff) / sigma_dsr
                dsr_neff  = float(stats.norm.cdf(z_neff))
                dsr_pass_neff = dsr_neff > dsr_thr
                dsr_pass_nom  = bool(tests_arc.get("dsr", {}).get("passed", False))

                # --- Tabella comparativa ---
                def _ico(p): return "✅" if p else "❌"
                def _bg(p): return "background-color:#E8F5E9" if p else "background-color:#FFEBEE"

                rows_neff = [
                    ("Bonferroni",
                     f"t={t_val:.3f}",
                     f"soglia N nom. = {float(tests_arc.get('bonferroni',{}).get('soglia',0)):.2f}",
                     f"{_ico(bon_pass_nom)} {'PASS' if bon_pass_nom else 'FAIL'}",
                     f"soglia Neff = {sb_neff:.2f}",
                     f"{_ico(bon_pass_neff)} {'PASS' if bon_pass_neff else 'FAIL'}"),
                    ("MinBTL",
                     f"anni IS = {years_val:.1f}",
                     f"MinBTL N nom. = {float(tests_arc.get('minbtl',{}).get('soglia',0)):.1f}",
                     f"{_ico(minbtl_pass_nom)} {'PASS' if minbtl_pass_nom else 'FAIL'}",
                     f"MinBTL Neff = {minbtl_neff:.1f}",
                     f"{_ico(minbtl_pass_neff)} {'PASS' if minbtl_pass_neff else 'FAIL'}"),
                    ("DSR",
                     f"DSR (Neff) = {dsr_neff:.3f}",
                     f"soglia = {dsr_thr}",
                     f"{_ico(dsr_pass_nom)} {'PASS' if dsr_pass_nom else 'FAIL'} (N nom.)",
                     f"soglia = {dsr_thr}",
                     f"{_ico(dsr_pass_neff)} {'PASS' if dsr_pass_neff else 'FAIL'}"),
                ]
                df_neff = pd.DataFrame(rows_neff, columns=[
                    "Test", "Valore", f"Soglia (N nom.={n_nom_saved})", f"Esito N nom.",
                    f"Soglia (Neff={neff_saved})", "Esito Neff ✦"
                ])
                st.dataframe(df_neff, use_container_width=True, hide_index=True)

                n_pass_neff = sum([bon_pass_neff, minbtl_pass_neff, dsr_pass_neff])
                # T-stat e EMSR non dipendono da N: prendi dall'archivio
                tstat_pass = bool(tests_arc.get("ttest", {}).get("passed", False))
                emsr_pass  = bool(tests_arc.get("emsr",  {}).get("passed", False))
                n_pass_tot_neff = n_pass_neff + tstat_pass + emsr_pass
                vmap_neff = {5:"VALIDA",4:"QUASI VALIDA",3:"PARZIALE",
                             2:"NON IDONEA",1:"NON IDONEA",0:"NON IDONEA"}
                verdetto_neff = vmap_neff.get(n_pass_tot_neff, "NON IDONEA")
                if n_pass_tot_neff == 5:
                    st.markdown(f'<div class="verdict-pass">✅ VERDETTO CON Neff: {verdetto_neff} — {n_pass_tot_neff}/5</div>',
                                unsafe_allow_html=True)
                elif n_pass_tot_neff >= 3:
                    st.markdown(f'<div class="verdict-partial">⚠️ VERDETTO CON Neff: {verdetto_neff} — {n_pass_tot_neff}/5</div>',
                                unsafe_allow_html=True)
                else:
                    st.markdown(f'<div class="verdict-fail">✗ VERDETTO CON Neff: {verdetto_neff} — {n_pass_tot_neff}/5</div>',
                                unsafe_allow_html=True)

                # Salva verdetto Neff nel session state per l'archivio
                st.session_state[CSCV_KEY]["verdetto_neff"]   = verdetto_neff
                st.session_state[CSCV_KEY]["n_pass_neff"]     = n_pass_tot_neff
                st.session_state[CSCV_KEY]["neff_bon_soglia"] = round(sb_neff, 3)
                st.session_state[CSCV_KEY]["neff_minbtl"]     = round(minbtl_neff, 2) if not np.isnan(minbtl_neff) else None
                st.session_state[CSCV_KEY]["neff_dsr"]        = round(dsr_neff, 3)

            else:
                st.info("ℹ️ Per il ricalcolo L1 con Neff, esegui prima il **Livello 1** "
                        "nel Tab 1 e assicurati che sia collegato a questa sessione CSCV.")

        interp = "\n".join(interp_lines)
        st.code(interp, language="")

        # ── Auto-salvataggio CSCV ────────────────────────
        st.markdown("---")
        _cscv_save_key = f"cscv_saved_{arc_id_saved}_{cscv_result.get('pbo','')}"
        if arc_id_saved and st.session_state.get("_l2_save_key") != _cscv_save_key:
            try:
                # Arricchisce il result con i campi Neff prima di salvare
                _res_to_save = dict(cscv_result)
                _res_to_save["neff"]          = saved.get("neff")
                _res_to_save["n_nominal"]     = saved.get("n_nominal")
                _res_to_save["verdetto_neff"] = saved.get("verdetto_neff")
                _res_to_save["n_pass_neff"]   = saved.get("n_pass_neff")
                ok = arc_add_cscv(arc_id_saved, _res_to_save, matrix_saved)
                if ok:
                    st.session_state["_l2_save_key"] = _cscv_save_key
                    st.success("✅ **Livello 2 (CSCV) salvato automaticamente** "
                               "→ vai al Tab 3 per il report completo L1+L2")
                else:
                    st.error("Errore nel salvataggio automatico CSCV.")
            except Exception as e:
                st.error(f"Errore: {e}")
        elif arc_id_saved:
            st.success("✅ **Livello 2 già in archivio** → vai al Tab 3 per il report L1+L2")
        else:
            st.markdown("""
            <div class="info-warn">
            💡 Esegui prima il <b>Livello 1</b> nel Tab 1 — il collegamento è automatico.
            </div>""", unsafe_allow_html=True)

        # ── Download ─────────────────────────────────────
        st.markdown("---")
        cd1, cd2, cd3 = st.columns(3)
        with cd1:
            st.download_button("💾 Scarica Report (.txt)",
                data=f"CSCV REPORT\n{'='*55}\n{interp}",
                file_name=f"cscv_{strategy_name.replace(' ','_')}.txt",
                mime="text/plain", key="dl_cscv_txt",
                use_container_width=True)
        with cd2:
            st.download_button("🖼️ Scarica Grafici (.png)",
                data=img_bytes,
                file_name=f"cscv_grafici_{strategy_name.replace(' ','_')}.png",
                mime="image/png", key="dl_cscv_png",
                use_container_width=True)
        with cd3:
            if DOCX_OK:
                CSCV_DOCX_KEY = "cscv_docx_bytes"
                if st.button("📄 Esporta Report Word (solo L2)",
                             key="btn_cscv_docx",
                             use_container_width=True,
                             help="Report Word con solo i dati CSCV. "
                                  "Per il report completo L1+L2 usa il Tab 3 Archivio."):
                    docx_b, docx_e = cscv_export_docx(
                        cscv_result, interp_lines,
                        strategy_name, img_bytes)
                    if docx_e:
                        st.error(f"Errore: {docx_e}")
                    else:
                        st.session_state[CSCV_DOCX_KEY] = {
                            "bytes": docx_b,
                            "fname": strategy_name.replace(" ","_"),
                        }
                if "cscv_docx_bytes" in st.session_state:
                    sd = st.session_state["cscv_docx_bytes"]
                    st.download_button(
                        "💾 Scarica documento Word",
                        data=sd["bytes"],
                        file_name=f"cscv_{sd['fname']}.docx",
                        mime="application/vnd.openxmlformats-officedocument"
                             ".wordprocessingml.document",
                        key="dl_cscv_docx",
                        use_container_width=True,
                    )


# ══════════════════════════════════════════════════════════════
# TAB 3 — ARCHIVIO
# ══════════════════════════════════════════════════════════════
with tab_a:
    st.markdown('<div class="section-header">📦 Archivio Strategie</div>',
                unsafe_allow_html=True)

    manifest = arc_manifest()
    st.caption(f"📁 Cartella: `{ARCHIVE_DIR.resolve()}`")

    if not manifest:
        st.info("L'archivio è vuoto. Carica un file nel Tab 1 per iniziare — il salvataggio è automatico.")
        st.stop()

    tot  = len(manifest)
    val  = sum(1 for m in manifest if m["verdetto"]=="VALIDA")
    has_c= sum(1 for m in manifest if m.get("has_cscv"))
    rc1,rc2,rc3 = st.columns(3)
    rc1.metric("Strategie archiviate", tot)
    rc2.metric("Validate (5/5)",       val)
    rc3.metric("Con CSCV",             has_c)
    st.markdown("---")

    # Filtri
    cf1,cf2,cf3 = st.columns(3)
    with cf1: filt_n = st.text_input("Filtra per nome","")
    with cf2: filt_v = st.selectbox("Verdetto",
        ["Tutti","VALIDA","QUASI VALIDA","PARZIALE","NON IDONEA"])
    with cf3: filt_c = st.selectbox("CSCV",
        ["Tutti","Con CSCV","Senza CSCV"])

    filtered = manifest
    if filt_n:    filtered = [m for m in filtered if filt_n.lower() in m["strategy_name"].lower()]
    if filt_v!="Tutti": filtered = [m for m in filtered if m["verdetto"]==filt_v]
    if filt_c=="Con CSCV":    filtered = [m for m in filtered if m.get("has_cscv")]
    elif filt_c=="Senza CSCV":filtered = [m for m in filtered if not m.get("has_cscv")]

    st.markdown(f"**{len(filtered)} strategie trovate**")

    for item in filtered:
        verd = item["verdetto"]
        icon = "✅" if verd=="VALIDA" else ("⚠️" if "PARZIALE" in verd else "❌")
        cscv_badge = " 🔄" if item.get("has_cscv") else ""
        pbo_str    = f" | PBO: {item['pbo']:.1%}" if item.get("pbo") is not None else ""
        dt_str     = item["archived_at"][:10]

        with st.expander(
            f"{icon} **{item['strategy_name']}** — {item['instrument']} "
            f"| {verd} ({item['n_passed']}/5){pbo_str}{cscv_badge} | {dt_str}"
        ):
            entry, err = arc_load(item["id"])
            if err: st.error(err); continue

            meta = entry["metadata"]
            if meta.get("note"): st.markdown(f"*{meta['note']}*")

            ea1,ea2,ea3,ea4 = st.columns(4)
            ea1.metric("Capitale",     f"${meta['initial_capital']:,}")
            ea2.metric("IS fino al",   meta["is_end_date"])
            ea3.metric("N combinaz.",  f"{meta['n_totale']:,}")
            ea4.metric("Archiviato",   meta["archived_at"][:10])

            if entry.get("morrone"):
                mor = entry["morrone"]
                m_is_a  = mor.get("metrics_is",{})
                m_oos_a = mor.get("metrics_oos",{})
                st.markdown("**Test statistici:**")
                test_cols = st.columns(5)
                for i,(tk,tv) in enumerate(mor["tests"].items()):
                    with test_cols[i]:
                        ico = "✅" if tv["passed"] else "❌"
                        st.markdown(
                            f'<div style="text-align:center;padding:0.5rem;'
                            f'background:{"#E8F5E9" if tv["passed"] else "#FFEBEE"};'
                            f'border-radius:8px;font-size:0.85rem;">'
                            f'{ico} <b>{tv["name"]}</b><br>'
                            f'{tv["value"]} / {tv["soglia"]}'
                            f'</div>', unsafe_allow_html=True)
                st.markdown("")
                mc1,mc2 = st.columns(2)
                with mc1:
                    st.markdown("**Metriche IS:**")
                    if m_is_a:
                        st.write({
                            "Sharpe":  round(float(m_is_a.get("sharpe",0)),3),
                            "CAGR":    f"{float(m_is_a.get('cagr',0)):.2%}",
                            "Max DD":  f"{float(m_is_a.get('max_dd',0)):.2%}",
                            "PF":      round(float(m_is_a.get("profit_factor",0)),2),
                            "Trade":   int(m_is_a.get("n_trades",0)),
                        })
                with mc2:
                    st.markdown("**Metriche OOS:**")
                    if m_oos_a:
                        st.write({
                            "Sharpe":  round(float(m_oos_a.get("sharpe",0)),3),
                            "CAGR":    f"{float(m_oos_a.get('cagr',0)):.2%}",
                            "Max DD":  f"{float(m_oos_a.get('max_dd',0)):.2%}",
                            "PF":      round(float(m_oos_a.get("profit_factor",0)),2),
                            "Trade":   int(m_oos_a.get("n_trades",0)),
                        })
                    else:
                        st.info("Nessun dato OOS.")

            if entry.get("cscv"):
                cr_a = entry["cscv"]
                st.markdown("---")
                neff_a   = cr_a.get("neff")
                n_nom_a  = cr_a.get("n_nominal")
                neff_str = f" | Neff=`{neff_a}`" if neff_a else ""
                verd_neff_a = cr_a.get("verdetto_neff")
                verd_neff_str = f" | Verdetto Neff: **{verd_neff_a}**" if verd_neff_a else ""
                st.markdown(f"**CSCV:** PBO=`{cr_a['pbo']:.1%}` | "
                            f"Prob.perdita=`{cr_a['prob_loss']:.1%}` | "
                            f"Config={cr_a['N_configs']}{neff_str} | "
                            f"Combo={cr_a['n_combos']:,}{verd_neff_str}")
                if st.button("📊 Mostra grafici CSCV",
                             key=f"cg_{item['id']}"):
                    fg = make_cscv_charts(cr_a)
                    st.pyplot(fg, use_container_width=True)
                    plt.close(fg)

            # Rielaborazione con parametri aggiornati
            if entry.get("trades") is not None:
                st.markdown("---")
                if st.button("🔄 Rielabora con parametri sidebar attuali",
                             key=f"rel_{item['id']}",
                             help="Ricarica i trade e riesegui i test con N, alpha, "
                                  "data IS e capitale della sidebar"):
                    tr2 = entry["trades"].copy()
                    if "return_pct" not in tr2.columns:
                        ec = initial_capital + tr2["profit"].cumsum()
                        ep = ec.shift(1).fillna(initial_capital)
                        tr2["equity"]     = ec
                        tr2["equity_prev"]= ep
                        tr2["return_pct"] = tr2["profit"]/ep
                    is_e2 = pd.Timestamp(is_end_date)
                    di2   = tr2[tr2["exit_datetime"]<=is_e2].copy().reset_index(drop=True)
                    if len(di2)>0:
                        di2["equity"]     = initial_capital+di2["profit"].cumsum()
                        di2["equity_prev"]= di2["equity"].shift(1).fillna(initial_capital)
                        di2["return_pct"] = di2["profit"]/di2["equity_prev"]
                        mi2 = compute_metrics(di2, initial_capital, "IS")
                        if mi2:
                            tr2_res,_,_ = run_all_tests(mi2,n_totale,alpha,dsr_threshold)
                            np2 = sum(1 for t in tr2_res.values() if t["passed"])
                            vmap3 = {5:"VALIDA",4:"QUASI VALIDA",3:"PARZIALE",
                                     2:"NON IDONEA",1:"NON IDONEA",0:"NON IDONEA"}
                            st.success(f"Rielaborazione completata: "
                                       f"**{vmap3.get(np2,'NON IDONEA')} — {np2}/5**")
                            for tv in tr2_res.values():
                                ico2 = "✅" if tv["passed"] else "❌"
                                st.markdown(f"{ico2} **{tv['name']}** — "
                                            f"{tv['value']} / {tv['soglia']}")

            st.markdown("---")
            if DOCX_OK:
                st.markdown("**📄 Esporta Report Word:**")
                has_cscv_e = item.get("has_cscv", False)
                fname_base = (meta.get("strategy_name","strategia")
                              .replace(" ","_").replace("/","_"))

                # 3 pulsanti chiari affiancati
                _ncols = 3 if has_cscv_e else 2
                _btn_cols = st.columns(_ncols + 1)  # +1 per elimina

                # Pulsante L1
                with _btn_cols[0]:
                    _key_l1 = f"docx_l1_{item['id']}"
                    if st.button("📄 Report L1", key=f"exp_l1_{item['id']}",
                                 use_container_width=True, help="Solo Livello 1 — Morrone"):
                        _b, _e = arc_export_docx(entry)
                        if _e: st.error(_e)
                        else: st.session_state[_key_l1] = {"bytes":_b, "fname":f"L1_{fname_base}"}
                    if _key_l1 in st.session_state:
                        st.download_button("💾 Scarica L1",
                            data=st.session_state[_key_l1]["bytes"],
                            file_name=f"{st.session_state[_key_l1]['fname']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"dl_l1_{item['id']}", use_container_width=True)

                # Pulsante L2 (solo se CSCV disponibile)
                if has_cscv_e:
                    with _btn_cols[1]:
                        _key_l2 = f"docx_l2_{item['id']}"
                        _cr_a = entry.get("cscv", {})
                        if st.button("📄 Report L2", key=f"exp_l2_{item['id']}",
                                     use_container_width=True, help="Solo Livello 2 — CSCV"):
                            # Genera grafico CSCV dall'archivio
                            _fig_tmp = make_cscv_charts(_cr_a)
                            _buf_tmp = io.BytesIO()
                            _fig_tmp.savefig(_buf_tmp, format="png", dpi=120,
                                             bbox_inches="tight", facecolor="#F8FAFC")
                            plt.close(_fig_tmp); _buf_tmp.seek(0)
                            _ilines = [
                                f"PBO: {_cr_a.get('pbo',0):.1%}",
                                f"Prob.perdita OOS: {_cr_a.get('prob_loss',0):.1%}",
                                f"N configurazioni: {_cr_a.get('N_configs','—')}",
                                f"N combinazioni: {int(_cr_a.get('n_combos',0)):,}",
                            ]
                            _b, _e = cscv_export_docx(
                                _cr_a, _ilines, meta.get("strategy_name",""), _buf_tmp.getvalue())
                            if _e: st.error(_e)
                            else: st.session_state[_key_l2] = {"bytes":_b, "fname":f"L2_{fname_base}"}
                        if _key_l2 in st.session_state:
                            st.download_button("💾 Scarica L2",
                                data=st.session_state[_key_l2]["bytes"],
                                file_name=f"{st.session_state[_key_l2]['fname']}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_l2_{item['id']}", use_container_width=True)

                # Pulsante L1+L2 (solo se CSCV disponibile)
                if has_cscv_e:
                    with _btn_cols[2]:
                        _key_both = f"docx_both_{item['id']}"
                        if st.button("📄 Report L1+L2", key=f"exp_both_{item['id']}",
                                     use_container_width=True, type="primary",
                                     help="Report completo con entrambi i livelli"):
                            _b, _e = combined_export_docx(entry)
                            if _e: st.error(_e)
                            else: st.session_state[_key_both] = {"bytes":_b, "fname":f"Completo_{fname_base}"}
                        if _key_both in st.session_state:
                            st.download_button("💾 Scarica L1+L2",
                                data=st.session_state[_key_both]["bytes"],
                                file_name=f"{st.session_state[_key_both]['fname']}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_both_{item['id']}", use_container_width=True)

                # Pulsante Elimina (sempre nell'ultima colonna)
                with _btn_cols[-1]:
                    if st.button("🗑️ Elimina", key=f"del_{item['id']}",
                                 type="secondary", use_container_width=True):
                        arc_delete(item["id"])
                        st.warning("Entry eliminata.")
                        st.rerun()
            else:
                st.warning("⚠️ python-docx non installato. Riavvia il .bat.")
                if st.button("🗑️ Elimina", key=f"del_{item['id']}",
                             type="secondary"):
                    arc_delete(item["id"])
                    st.warning("Entry eliminata.")
                    st.rerun()

    # Export riepilogo
    st.markdown("---")
    st.markdown("### 📤 Export riepilogo archivio")
    ce1,ce2 = st.columns(2)
    with ce1:
        if manifest:
            rows_e = [{
                "Strategia":     m["strategy_name"],
                "Strumento":     m["instrument"],
                "Archiviato":    m["archived_at"][:10],
                "Verdetto":      m["verdetto"],
                "Test superati": m["n_passed"],
                "Ha CSCV":       "Si" if m.get("has_cscv") else "No",
                "PBO":           f"{m['pbo']:.1%}" if m.get("pbo") is not None else "-",
                "Note":          m.get("note",""),
            } for m in manifest]
            csv_e = pd.DataFrame(rows_e).to_csv(index=False, sep=";", decimal=",")
            st.download_button("📊 Scarica riepilogo (.csv)",
                data=csv_e, file_name="archivio_strategie.csv",
                mime="text/csv")
    with ce2:
        st.markdown("""
        <div class="info-box">
        La cartella <b>validator_archive/</b> contiene tutti i dati
        in formato standard (JSON + CSV). Copiala su un altro PC
        per trasferire l'intero archivio.
        </div>""", unsafe_allow_html=True)
